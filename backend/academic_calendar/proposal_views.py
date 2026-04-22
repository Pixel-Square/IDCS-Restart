"""
API views for the multi-stage Event Proposal approval workflow.

Workflow:  Staff → Branding → HOD → HAA → Done (notifications).


Endpoints mounted at /api/academic-calendar/proposals/…
"""
from __future__ import annotations

import base64
import re
import logging
import mimetypes
import uuid
from io import BytesIO
from datetime import date, datetime
from typing import Any, Dict, Optional
from urllib.parse import unquote, urlparse
from urllib.request import Request as UrlRequest, urlopen
from zoneinfo import ZoneInfo

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.core.mail import send_mail
from django.db import transaction
from django.http import Http404, HttpResponse
from django.utils import timezone
from django.views.decorators.http import require_http_methods

from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework_simplejwt.authentication import JWTAuthentication

from academics.models import AcademicYear, Department, DepartmentRole, StaffProfile
from accounts.models import UserNotification, UserRole
from accounts.utils import get_user_permissions

from .models import EventProposal

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
except Exception:  # pragma: no cover
    Document = None
    WD_SECTION = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None

try:
    from PIL import Image as PILImage
except Exception:  # pragma: no cover
    PILImage = None

logger = logging.getLogger(__name__)

IST_TIMEZONE = ZoneInfo('Asia/Kolkata')


def _is_branding_workflow_user(perms: set[str]) -> bool:
    return 'branding.access' in perms and 'events.branding_review' in perms


def _has_branding_read_access(perms: set[str]) -> bool:
    return (
        'events.branding_review' in perms
        or 'branding.list_posters' in perms
        or 'branding.access' in perms
    )


def _jwt_authenticate_request(request):
    """Return authenticated user from Authorization header, or None."""
    try:
        res = JWTAuthentication().authenticate(request)
        if res:
            user, _ = res
            return user
    except Exception:
        return None
    return None


def _jwt_authenticate_query_param_token(request):
    """Return authenticated user from ?token=, or None."""
    qs_token = request.GET.get('token', '').strip()
    if not qs_token:
        return None
    from django.http import HttpRequest as _Req

    fake_req = _Req()
    fake_req.META['HTTP_AUTHORIZATION'] = f'Bearer {qs_token}'
    try:
        res = JWTAuthentication().authenticate(fake_req)
        if res:
            user, _ = res
            return user
    except Exception:
        return None
    return None


def _can_user_view_proposal(user, perms, proposal: EventProposal) -> bool:
    if not user or not getattr(user, 'is_authenticated', False):
        return False
    if not proposal:
        return False

    if getattr(user, 'is_superuser', False):
        return True

    # Creator can always view their own proposal media, even if permission
    # assignments changed after the proposal was created.
    if proposal.created_by_id == user.id:
        return True

    # Branding: can view workflow proposals
    if _has_branding_read_access(perms):
        return True

    # HAA reviewers can view proposal media in the workflow.
    if 'events.haa_approve' in perms:
        return True

    # HOD reviewers can view proposal media in the workflow.
    # Do not enforce department gate in download endpoint, because HOD-stage
    # visibility is already constrained by list/detail workflows and this avoids
    # false 403s caused by legacy/missing mappings.
    if 'events.hod_approve' in perms:
        return True

    # Fallback: allow core workflow roles even if role-permission mappings are
    # temporarily inconsistent in DB.
    try:
        has_workflow_role = UserRole.objects.filter(
            user=user,
            role__name__in=['HOD', 'HAA', 'BRANDING'],
        ).exists()
        if has_workflow_role:
            return True
    except Exception:
        pass

    return False


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _user_display_name(user) -> str:
    full = user.get_full_name()
    return full if full.strip() else user.username


def _format_ist_datetime(value: datetime | None) -> str:
    if not value:
        return ''
    dt = value
    if timezone.is_naive(dt):
        dt = timezone.make_aware(dt, timezone.get_current_timezone())
    dt = timezone.localtime(dt, IST_TIMEZONE)
    return dt.strftime('%d-%m-%Y %I:%M %p IST')


def _get_active_academic_year():
    return AcademicYear.objects.filter(is_active=True).first()


def _get_hod_for_department(department: Department) -> Optional[Any]:
    """Return the User who is the active HOD for *department* in the current year."""
    ay = _get_active_academic_year()
    if not ay or not department:
        return None
    dr = (
        DepartmentRole.objects
        .filter(department=department, role='HOD', is_active=True, academic_year=ay)
        .select_related('staff__user')
        .first()
    )
    return dr.staff.user if dr else None


def _get_haa_user():
    """Return the User with HAA role (fac id 3142016 per requirements)."""
    try:
        return (
            UserRole.objects
            .filter(role__name='HAA')
            .select_related('user')
            .first()
            .user
        )
    except (AttributeError, Exception):
        return None


def _iso(val):
    """Safely convert date/datetime/string to ISO string. Returns None if falsy."""
    if not val:
        return None
    if hasattr(val, 'isoformat'):
        return val.isoformat()
    return str(val)


def _resolve_local_media_rel_path_from_url(raw_url: str) -> str:
    """Resolve MEDIA_URL-backed URL/path into a storage-relative path."""
    candidate = str(raw_url or '').strip()
    if not candidate:
        return ''

    media_url_local = str(getattr(settings, 'MEDIA_URL', '/media/') or '/media/')
    parsed_local = urlparse(candidate)
    parsed_path_local = unquote(parsed_local.path or '')
    if candidate.startswith(media_url_local):
        parsed_path_local = candidate
    if not parsed_path_local.startswith(media_url_local):
        return ''

    rel = parsed_path_local[len(media_url_local):].lstrip('/').replace('\\', '/')
    return rel


def _serialize_proposal(p: EventProposal) -> Dict:
    poster_url_text = str(p.poster_url or '').strip()
    legacy_rel_path = _resolve_local_media_rel_path_from_url(poster_url_text)
    has_review_stage_poster = (
        bool(poster_url_text)
        and p.status in {
            EventProposal.Status.FORWARDED_TO_HOD,
            EventProposal.Status.HOD_APPROVED,
            EventProposal.Status.FORWARDED_TO_HAA,
            EventProposal.Status.HAA_APPROVED,
        }
    )
    has_final_poster = (
        bool(str(p.final_poster_path or '').strip())
        or legacy_rel_path.startswith('proposal-posters/')
        or has_review_stage_poster
    )

    return {
        'id': str(p.id),
        'title': p.title,
        'department_name': p.department_name,
        'event_type': p.event_type,
        'start_date': _iso(p.start_date),
        'end_date': _iso(p.end_date),
        'venue': p.venue,
        'mode': p.mode,
        'expert_category': p.expert_category,
        'is_repeated': p.is_repeated,
        'participants': p.participants,
        'coordinator_name': p.coordinator_name,
        'co_coordinator_name': p.co_coordinator_name,
        'chief_guest_name': p.chief_guest_name,
        'chief_guest_designation': p.chief_guest_designation,
        'chief_guest_affiliation': p.chief_guest_affiliation,
        'poster_url': p.poster_url,
        'poster_data_url': p.poster_data_url,
        'has_final_poster': has_final_poster,
        'proposal_doc_url': p.proposal_doc_url,
        'proposal_doc_name': p.proposal_doc_name,
        'canva_design_id': p.canva_design_id,
        'canva_edit_url': p.canva_edit_url,
        'status': p.status,
        'status_display': p.get_status_display(),
        'created_by_name': p.created_by_name,
        'branding_reviewed_by_name': p.branding_reviewed_by_name,
        'branding_reviewed_at': _iso(p.branding_reviewed_at),
        'branding_note': p.branding_note,
        'hod_approved_by_name': p.hod_approved_by_name,
        'hod_approved_at': _iso(p.hod_approved_at),
        'hod_note': p.hod_note,
        'haa_approved_by_name': p.haa_approved_by_name,
        'haa_approved_at': _iso(p.haa_approved_at),
        'haa_note': p.haa_note,
        'rejection_reason': p.rejection_reason,
        'rejected_at': _iso(p.rejected_at),
        'created_at': _iso(p.created_at),
        'updated_at': _iso(p.updated_at),
        'proposal_data': p.proposal_data,
    }


def _first_non_empty(*values: Any) -> str:
    for value in values:
        text = str(value or '').strip()
        if text:
            return text
    return ''


def _build_doc_payload_for_proposal(proposal: EventProposal, base_payload: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    payload: Dict[str, Any] = dict(base_payload or {})

    from_name = _first_non_empty(
        payload.get('from_name'),
        payload.get('from'),
        payload.get('proposalFrom'),
        proposal.created_by_name,
    )
    if from_name:
        payload['from_name'] = from_name

    department_name = _first_non_empty(
        payload.get('organizer_department_doc'),
        payload.get('organizer_department_raw'),
        payload.get('organizer_department'),
        payload.get('department'),
        proposal.department_name,
        getattr(proposal.department, 'name', ''),
    )
    if department_name:
        payload['organizer_department_doc'] = department_name
        payload['organizer_department_raw'] = department_name
        payload['organizer_department'] = department_name
        payload['department'] = department_name

    payload['event_type'] = _first_non_empty(payload.get('event_type'), proposal.event_type)
    payload['event_title'] = _first_non_empty(payload.get('event_title'), payload.get('event_name'), payload.get('title'), proposal.title)
    payload['participants'] = _first_non_empty(payload.get('participants'), proposal.participants)
    payload['coordinator'] = _first_non_empty(payload.get('coordinator'), payload.get('committee_member_1_name'), proposal.coordinator_name)
    payload['co_coordinator'] = _first_non_empty(payload.get('co_coordinator'), payload.get('committee_member_2_name'), proposal.co_coordinator_name)
    payload['resource_person'] = _first_non_empty(payload.get('resource_person'), payload.get('chief_guest_name'), proposal.chief_guest_name)
    payload['designation'] = _first_non_empty(payload.get('designation'), payload.get('chief_guest_position'), proposal.chief_guest_designation)
    payload['resource_person_affiliation'] = _first_non_empty(
        payload.get('resource_person_affiliation'),
        payload.get('chief_guest_company'),
        proposal.chief_guest_affiliation,
    )

    if proposal.start_date and not str(payload.get('from_date') or '').strip():
        payload['from_date'] = proposal.start_date.strftime('%d %B %Y')
    if proposal.end_date and not str(payload.get('to_date') or '').strip():
        payload['to_date'] = proposal.end_date.strftime('%d %B %Y')

    if not str(payload.get('event_date') or '').strip() and proposal.start_date:
        if proposal.end_date and proposal.end_date != proposal.start_date:
            payload['event_date'] = f"From {proposal.start_date.strftime('%d %B %Y')} to {proposal.end_date.strftime('%d %B %Y')}"
        else:
            payload['event_date'] = proposal.start_date.strftime('%d %B %Y')

    return payload


def _create_notification(user, title, message, link='', data=None):
    """Create an in-app UserNotification."""
    if user is None:
        return
    UserNotification.objects.create(
        user=user,
        title=title,
        message=message,
        link=link,
        data=data or {},
    )


@require_http_methods(['GET'])
def proposal_poster_download(request, proposal_id):
    """Serve a proposal poster image.

    Accepts auth either as:
      • Authorization: Bearer <token>
      • ?token=<access_token>

    Returns an inline image response if stored as a data URL.
    """
    user = getattr(request, 'user', None)
    if not user or not getattr(user, 'is_authenticated', False):
        user = _jwt_authenticate_request(request)
    if not user or not getattr(user, 'is_authenticated', False):
        user = _jwt_authenticate_query_param_token(request)

    if not user or not getattr(user, 'is_authenticated', False):
        return HttpResponse(
            '{"detail":"Authentication credentials were not provided."}',
            status=401,
            content_type='application/json',
        )

    proposal = (
        EventProposal.objects
        .filter(id=proposal_id)
        .select_related('department', 'created_by')
        .first()
    )
    if not proposal:
        raise Http404('Poster not found.')

    perms = get_user_permissions(user)
    if not _can_user_view_proposal(user, perms, proposal):
        return HttpResponse('{"detail":"Forbidden"}', status=403, content_type='application/json')

    def _stream_storage_path(storage_path: str) -> HttpResponse | None:
        if not storage_path or not default_storage.exists(storage_path):
            return None
        try:
            with default_storage.open(storage_path, 'rb') as media_file:
                content = media_file.read()
            mime, _ = mimetypes.guess_type(storage_path)
            response = HttpResponse(content, content_type=mime or 'application/octet-stream')
            response['Content-Disposition'] = f'inline; filename="{storage_path.rsplit("/", 1)[-1]}"'
            response['Cache-Control'] = 'no-store'
            return response
        except Exception as exc:
            logger.warning('Failed to stream local poster file %s: %s', storage_path, exc)
            return None

    final_only = str(request.GET.get('final_only', '') or '').strip().lower() in {'1', 'true', 'yes'}
    if final_only:
        final_response = _stream_storage_path(str(proposal.final_poster_path or '').strip())
        if final_response is not None:
            return final_response

        # Legacy fallback: only trust proposal-posters/* as Branding final uploads.
        legacy_rel_path = _resolve_local_media_rel_path_from_url(str(proposal.poster_url or '').strip())
        if legacy_rel_path.startswith('proposal-posters/'):
            legacy_response = _stream_storage_path(legacy_rel_path)
            if legacy_response is not None:
                return legacy_response

        poster_url_text = str(proposal.poster_url or '').strip()
        is_review_stage = proposal.status in {
            EventProposal.Status.FORWARDED_TO_HOD,
            EventProposal.Status.HOD_APPROVED,
            EventProposal.Status.FORWARDED_TO_HAA,
            EventProposal.Status.HAA_APPROVED,
        }
        if is_review_stage and (poster_url_text.startswith('http://') or poster_url_text.startswith('https://') or poster_url_text.startswith('/')):
            response = HttpResponse(status=302)
            response['Location'] = poster_url_text
            response['Cache-Control'] = 'no-store'
            return response

        raise Http404('Final poster not found.')

    src = (proposal.poster_data_url or proposal.poster_url or '').strip()
    if not src:
        raise Http404('Poster not found.')

    if src.startswith('data:image/') and 'base64,' in src:
        header, b64 = src.split('base64,', 1)
        mime = header[5:].split(';', 1)[0] if header.startswith('data:') else 'image/png'
        try:
            raw = base64.b64decode(b64, validate=False)
        except Exception:
            raise Http404('Poster not found.')

        ext = 'png'
        if '/' in mime:
            ext = mime.split('/', 1)[1] or 'png'

        response = HttpResponse(raw, content_type=mime)
        response['Content-Disposition'] = f'inline; filename="poster.{ext}"'
        response['Cache-Control'] = 'no-store'
        return response

    # If src points to local MEDIA_URL, stream from storage directly.
    final_response = _stream_storage_path(str(proposal.final_poster_path or '').strip())
    if final_response is not None:
        return final_response

    media_url = str(getattr(settings, 'MEDIA_URL', '/media/') or '/media/')
    parsed = urlparse(src)
    parsed_path = unquote(parsed.path or '')
    request_host = str(request.get_host() or '').split(':', 1)[0].strip().lower()
    allowed_hosts = {
        'localhost',
        '127.0.0.1',
        '0.0.0.0',
    }
    if request_host:
        allowed_hosts.add(request_host)
    for host in (getattr(settings, 'ALLOWED_HOSTS', []) or []):
        host_text = str(host or '').strip().lower()
        if not host_text:
            continue
        if host_text.startswith('.'):
            host_text = host_text[1:]
        if host_text:
            allowed_hosts.add(host_text)
    if src.startswith(media_url):
        parsed_path = src
    parsed_host = str(parsed.hostname or '').strip().lower()
    if parsed_path.startswith(media_url) and (not parsed_host or parsed_host in allowed_hosts):
        rel_path = parsed_path[len(media_url):].lstrip('/').replace('\\', '/')
        if rel_path and default_storage.exists(rel_path):
            media_response = _stream_storage_path(rel_path)
            if media_response is not None:
                return media_response

    # Fallback: redirect to URL (if present)
    if src.startswith('http://') or src.startswith('https://') or src.startswith('/'):
        response = HttpResponse(status=302)
        response['Location'] = src
        response['Cache-Control'] = 'no-store'
        return response

    raise Http404('Poster not found.')


def _extract_doc_storage_path_from_url(raw_url: str) -> str:
    """Extract proposal-docs/<id>/event-proposal-format.docx from URL/path."""
    text = str(raw_url or '').strip()
    if not text:
        return ''
    parsed = urlparse(text)
    path = unquote(parsed.path or '')
    match = re.search(r'/proposal-docs/([a-zA-Z0-9_-]+)/event-proposal-format\.docx$', path)
    if not match:
        return ''
    return f'proposal-docs/{match.group(1)}/event-proposal-format.docx'


def _load_proposal_doc_bytes(proposal: EventProposal, request=None) -> bytes:
    storage_path = _extract_doc_storage_path_from_url(proposal.proposal_doc_url)
    if storage_path and default_storage.exists(storage_path):
        with default_storage.open(storage_path, 'rb') as fh:
            return fh.read()

    # Fallback: regenerate the latest proposal form if link is stale/missing.
    if proposal.status == EventProposal.Status.HAA_APPROVED:
        _regenerate_approved_doc(proposal, request=request)
    elif proposal.status in {EventProposal.Status.HOD_APPROVED, EventProposal.Status.FORWARDED_TO_HAA}:
        _regenerate_doc_after_hod(proposal, request=request)
    else:
        _regenerate_doc_after_hod(proposal, request=request)

    proposal.refresh_from_db(fields=['proposal_doc_url'])
    storage_path = _extract_doc_storage_path_from_url(proposal.proposal_doc_url)
    if storage_path and default_storage.exists(storage_path):
        with default_storage.open(storage_path, 'rb') as fh:
            return fh.read()

    raise Http404('Proposal document not found.')


def _regenerate_proposal_doc_for_download(proposal: EventProposal, request=None) -> None:
    """Refresh proposal DOCX so final-download always uses latest template/layout."""
    try:
        if proposal.status == EventProposal.Status.HAA_APPROVED:
            _regenerate_approved_doc(proposal, request=request)
        else:
            _regenerate_doc_after_hod(proposal, request=request)
    except Exception as exc:
        logger.warning('DOCX refresh before final-download failed for %s: %s', proposal.id, exc)


def _load_final_poster_bytes(proposal: EventProposal) -> bytes:
    final_path = str(proposal.final_poster_path or '').strip()
    if final_path and default_storage.exists(final_path):
        with default_storage.open(final_path, 'rb') as fh:
            return fh.read()

    poster_data_url = str(proposal.poster_data_url or '').strip()
    if poster_data_url.startswith('data:image/') and 'base64,' in poster_data_url:
        _, b64 = poster_data_url.split('base64,', 1)
        try:
            return base64.b64decode(b64, validate=False)
        except Exception:
            pass

    poster_url = str(proposal.poster_url or '').strip()
    rel_from_url = _resolve_local_media_rel_path_from_url(poster_url)
    if rel_from_url and default_storage.exists(rel_from_url):
        with default_storage.open(rel_from_url, 'rb') as fh:
            return fh.read()

    if poster_url.startswith('http://') or poster_url.startswith('https://'):
        try:
            req = UrlRequest(poster_url, headers={'User-Agent': 'IDCS-Backend/1.0'})
            with urlopen(req, timeout=15) as resp:
                return resp.read()
        except Exception:
            pass

    raise Http404('Final poster not found.')


def _build_docx_with_embedded_poster(doc_bytes: bytes, poster_bytes: bytes) -> bytes:
    if Document is None or Inches is None:
        raise RuntimeError('python-docx is not installed on the server.')

    doc = Document(BytesIO(doc_bytes))

    # Put poster on its own section so we can strip inherited template headers.
    if WD_SECTION is not None:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        doc.add_page_break()
        section = doc.sections[-1]

    # Remove template logos/headers from the poster page.
    for attr in ('header', 'first_page_header', 'even_page_header', 'footer', 'first_page_footer', 'even_page_footer'):
        part = getattr(section, attr, None)
        if part is None:
            continue
        try:
            part.is_linked_to_previous = False
        except Exception:
            pass
        try:
            for p in part.paragraphs:
                p.text = ''
        except Exception:
            pass

    # Full-bleed section for print-like output.
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    try:
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0)
    except Exception:
        pass

    image_para = doc.add_paragraph()
    image_para.paragraph_format.space_before = 0
    image_para.paragraph_format.space_after = 0
    try:
        image_para.paragraph_format.line_spacing = 1
    except Exception:
        pass
    if WD_ALIGN_PARAGRAPH is not None:
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_width = section.page_width
    page_height = section.page_height

    # Word can clip images that sit exactly on page boundaries, so keep a tiny
    # safety buffer (especially at the bottom) while still looking full-page.
    safe_left_right = Inches(0.03)
    safe_top = Inches(0.02)
    safe_bottom = Inches(0.10)
    max_width = max(1, int(page_width - (2 * safe_left_right)))
    max_height = max(1, int(page_height - safe_top - safe_bottom))

    fit_width = max_width
    fit_height = max_height
    if PILImage is not None:
        try:
            with PILImage.open(BytesIO(poster_bytes)) as img:
                img_width_px, img_height_px = img.size
            if img_width_px > 0 and img_height_px > 0:
                # Scale to fit whole image inside page (no cropping).
                width_scale = float(max_width) / float(img_width_px)
                height_scale = float(max_height) / float(img_height_px)
                scale = min(width_scale, height_scale)
                fit_width = max(1, int(img_width_px * scale))
                fit_height = max(1, int(img_height_px * scale))

                top_pad = int(safe_top + max(0, int((max_height - fit_height) / 2)))
                image_para.paragraph_format.space_before = top_pad
        except Exception:
            fit_width = max_width
            fit_height = max_height

    image_run = image_para.add_run()
    image_run.add_picture(BytesIO(poster_bytes), width=fit_width, height=fit_height)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


@require_http_methods(['GET'])
def proposal_final_doc_download(request, proposal_id):
    """Download one printable DOCX containing the event form + final poster."""
    user = getattr(request, 'user', None)
    if not user or not getattr(user, 'is_authenticated', False):
        user = _jwt_authenticate_request(request)
    if not user or not getattr(user, 'is_authenticated', False):
        user = _jwt_authenticate_query_param_token(request)

    if not user or not getattr(user, 'is_authenticated', False):
        return HttpResponse(
            '{"detail":"Authentication credentials were not provided."}',
            status=401,
            content_type='application/json',
        )

    proposal = (
        EventProposal.objects
        .filter(id=proposal_id)
        .select_related('department', 'created_by')
        .first()
    )
    if not proposal:
        raise Http404('Proposal not found.')

    perms = get_user_permissions(user)
    if not _can_user_view_proposal(user, perms, proposal):
        return HttpResponse('{"detail":"Forbidden"}', status=403, content_type='application/json')

    try:
        # Always refresh proposal DOCX so alignment/template fixes are visible
        # in every final-download without waiting for status transitions.
        _regenerate_proposal_doc_for_download(proposal, request=request)
        proposal.refresh_from_db(fields=['proposal_doc_url', 'proposal_doc_name'])

        doc_bytes = _load_proposal_doc_bytes(proposal, request=request)
        poster_bytes = _load_final_poster_bytes(proposal)
        merged = _build_docx_with_embedded_poster(doc_bytes, poster_bytes)
    except Http404:
        raise
    except Exception as exc:
        logger.warning('Failed to generate final combined DOCX for proposal %s: %s', proposal.id, exc)
        return HttpResponse(
            '{"detail":"Could not generate final combined document."}',
            status=500,
            content_type='application/json',
        )

    safe_title = re.sub(r'[^A-Za-z0-9_-]+', '_', str(proposal.title or 'event').strip()).strip('_') or 'event'
    filename = f'{safe_title}_Final_Approval_With_Poster.docx'
    response = HttpResponse(
        merged,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response['Cache-Control'] = 'no-store'
    return response


# ──────────────────────────────────────────────────────────────────────────────
# LIST / CREATE proposals
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def proposals_list_create(request):
    """
    GET  - list proposals visible to the calling user:
           • Staff:    own proposals
           • Branding: proposals with status forwarded_to_branding
           • HOD:      proposals forwarded_to_hod for their department
           • HAA:      proposals forwarded_to_haa
    POST - create a new proposal (status=forwarded_to_branding).
    """
    user = request.user
    perms = get_user_permissions(user)

    if request.method == 'GET':
        return _list_proposals(user, perms, request)
    else:
        return _create_proposal(user, perms, request)


def _list_proposals(user, perms, request):
    status_filter = request.query_params.get('status', '')
    mine_only = request.query_params.get('mine', '') in ('1', 'true', 'True')
    qs = EventProposal.objects.all()

    # If ?mine=1, always show only the user's own proposals regardless of role
    if mine_only:
        qs = qs.filter(created_by=user)
    elif 'events.haa_approve' in perms:
        # HAA sees proposals forwarded to them + final approved
        qs = qs.filter(status__in=[
            EventProposal.Status.FORWARDED_TO_HAA,
            EventProposal.Status.HAA_APPROVED,
            EventProposal.Status.REJECTED,
        ])
    elif 'events.hod_approve' in perms:
        # HOD sees proposals for ALL departments they manage
        try:
            hod_depts = _get_user_hod_departments(user)
        except Exception:
            hod_depts = []
        hod_statuses = [
            EventProposal.Status.FORWARDED_TO_HOD,
            EventProposal.Status.HOD_APPROVED,
            EventProposal.Status.FORWARDED_TO_HAA,
            EventProposal.Status.HAA_APPROVED,
            EventProposal.Status.REJECTED,
        ]
        if hod_depts:
            qs = qs.filter(department__in=hod_depts, status__in=hod_statuses)
        else:
            # HOD has no department mapping at all — show all HOD-stage proposals
            qs = qs.filter(status__in=hod_statuses)
    elif _has_branding_read_access(perms):
        # Branding sees forwarded proposals
        qs = qs.filter(status__in=[
            EventProposal.Status.FORWARDED_TO_BRANDING,
            EventProposal.Status.FORWARDED_TO_HOD,
            EventProposal.Status.HOD_APPROVED,
            EventProposal.Status.FORWARDED_TO_HAA,
            EventProposal.Status.HAA_APPROVED,
            EventProposal.Status.REJECTED,
        ])
    elif 'events.create_proposal' in perms:
        # Staff sees own proposals
        qs = qs.filter(created_by=user)
    else:
        qs = qs.none()

    # Optional extra status filter
    if status_filter:
        qs = qs.filter(status=status_filter)

    proposals = list(qs.order_by('-created_at')[:100])
    return Response([_serialize_proposal(p) for p in proposals])


def _get_user_hod_departments(user) -> list:
    """Return ALL Department objects the user is currently an active HOD of.

    Primary source: DepartmentRole (role='HOD', is_active=True, current AY).
    Fallback: StaffProfile.department (covers HODs not yet entered in DeptRole).
    Returns an empty list if nothing found.
    """
    depts: list = []
    try:
        ay = _get_active_academic_year()
        if ay:
            dept_ids_seen: set = set()
            for dr in (
                DepartmentRole.objects
                .filter(staff__user=user, role='HOD', is_active=True, academic_year=ay)
                .select_related('department')
            ):
                if dr.department and dr.department.id not in dept_ids_seen:
                    depts.append(dr.department)
                    dept_ids_seen.add(dr.department.id)
    except Exception:
        pass

    if not depts:
        # Fallback: staff profile department
        try:
            sp = user.staff_profile
            d = sp.get_current_department() if hasattr(sp, 'get_current_department') else getattr(sp, 'department', None)
            if d:
                depts.append(d)
        except Exception:
            pass

    return depts


def _get_user_department(user):
    """Return the first Department the user is HOD of (kept for backwards compat)."""
    depts = _get_user_hod_departments(user)
    return depts[0] if depts else None


def _create_proposal(user, perms, request):
    if 'events.create_proposal' not in perms:
        return Response({'error': 'Missing events.create_proposal permission'},
                        status=status.HTTP_403_FORBIDDEN)

    data = request.data
    title = data.get('title', '').strip()
    if not title:
        return Response({'error': 'title is required'}, status=status.HTTP_400_BAD_REQUEST)

    start_date = data.get('start_date')
    if not start_date:
        return Response({'error': 'start_date is required'}, status=status.HTTP_400_BAD_REQUEST)

    # Resolve department — from explicit form data first, then creator's own profile
    dept = None
    dept_name = data.get('department_name', '')
    dept_id = data.get('department_id')
    if dept_id:
        dept = Department.objects.filter(id=dept_id).first()
    if not dept and dept_name:
        dept = Department.objects.filter(name__iexact=dept_name.strip()).first()
        if not dept:
            dept = Department.objects.filter(code__iexact=dept_name.strip()).first()
    # Fallback: resolve from the creator's StaffProfile so the proposal
    # is always tagged to the correct department even when the form omits it.
    if not dept:
        try:
            sp = user.staff_profile
            dept = (
                sp.get_current_department()
                if hasattr(sp, 'get_current_department')
                else getattr(sp, 'department', None)
            )
        except Exception:
            pass
    # Use department name from resolved dept object if form didn't supply one
    if dept and not dept_name:
        dept_name = dept.name

    proposal_data = data.get('proposal_data', {})
    if not isinstance(proposal_data, dict):
        proposal_data = {}

    try:
        proposal = EventProposal.objects.create(
            title=title,
            department=dept,
            department_name=dept_name or (dept.name if dept else ''),
            event_type=data.get('event_type', ''),
            start_date=start_date,
            end_date=data.get('end_date') or start_date,
            venue=data.get('venue', ''),
            mode=data.get('mode', ''),
            expert_category=data.get('expert_category', ''),
            is_repeated=bool(data.get('is_repeated', False)),
            participants=data.get('participants', ''),
            coordinator_name=data.get('coordinator_name', ''),
            co_coordinator_name=data.get('co_coordinator_name', ''),
            chief_guest_name=data.get('chief_guest_name', ''),
            chief_guest_designation=data.get('chief_guest_designation', ''),
            chief_guest_affiliation=data.get('chief_guest_affiliation', ''),
            proposal_data=proposal_data,
            poster_url=data.get('poster_url', ''),
            poster_data_url=data.get('poster_data_url', ''),
            proposal_doc_url=data.get('proposal_doc_url', ''),
            proposal_doc_name=data.get('proposal_doc_name', ''),
            canva_design_id=data.get('canva_design_id', ''),
            canva_edit_url=data.get('canva_edit_url', ''),
            status=EventProposal.Status.FORWARDED_TO_BRANDING,
            created_by=user,
            created_by_name=_user_display_name(user),
        )
    except Exception as exc:
        import traceback
        traceback.print_exc()
        return Response({'error': f'Failed to save proposal: {exc}'}, status=status.HTTP_400_BAD_REQUEST)

    # Refresh so auto_now/auto_now_add fields are proper Python objects
    proposal.refresh_from_db()

    normalized_doc_payload = _build_doc_payload_for_proposal(proposal, proposal.proposal_data)
    if normalized_doc_payload != (proposal.proposal_data or {}):
        proposal.proposal_data = normalized_doc_payload
        proposal.save(update_fields=['proposal_data'])

    # Notify branding users
    _notify_role_users(
        'BRANDING',
        f'New Event Proposal: {title}',
        f'{_user_display_name(user)} has submitted an event proposal "{title}" for review.',
        link=f'/branding/list-posters',
        data={'proposal_id': str(proposal.id)},
    )

    return Response(_serialize_proposal(proposal), status=status.HTTP_201_CREATED)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def proposals_delete_all(request):
    """
    Delete all EventProposal records from the approval workflow.
    Restricted to Branding users only.
    """
    perms = get_user_permissions(request.user)
    has_dedicated_delete_perm = 'events.bulk_delete_proposals' in perms
    if not has_dedicated_delete_perm and not _is_branding_workflow_user(perms):
        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

    with transaction.atomic():
        deleted_count, _ = EventProposal.objects.all().delete()

    return Response({'ok': True, 'deleted_count': deleted_count})


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def branding_upload_final_poster(request, proposal_id):
    """Upload final refined poster for a proposal (Branding only)."""
    perms = get_user_permissions(request.user)
    if not _has_branding_read_access(perms):
        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

    try:
        proposal = EventProposal.objects.get(id=proposal_id)
    except EventProposal.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

    if proposal.status in (EventProposal.Status.HAA_APPROVED, EventProposal.Status.REJECTED):
        return Response({'error': 'Cannot update poster for finalized proposals'}, status=status.HTTP_400_BAD_REQUEST)

    uploaded_file = request.FILES.get('poster')
    if not uploaded_file:
        return Response({'error': 'poster file is required'}, status=status.HTTP_400_BAD_REQUEST)

    content_type = str(getattr(uploaded_file, 'content_type', '') or '').lower()
    if content_type and not content_type.startswith('image/'):
        return Response({'error': 'Only image files are allowed'}, status=status.HTTP_400_BAD_REQUEST)

    ext = 'png'
    original_name = str(getattr(uploaded_file, 'name', '') or '').lower()
    if '.' in original_name:
        ext = original_name.rsplit('.', 1)[1].strip() or 'png'

    save_path = default_storage.save(
        f'proposal-posters/{proposal.id}/{uuid.uuid4().hex}.{ext}',
        ContentFile(uploaded_file.read()),
    )

    media_url = str(getattr(settings, 'MEDIA_URL', '/media/') or '/media/')
    proposal.poster_url = request.build_absolute_uri(f'{media_url}{save_path}')
    proposal.poster_data_url = ''
    proposal.final_poster_path = save_path
    proposal.save(update_fields=['poster_url', 'poster_data_url', 'final_poster_path', 'updated_at'])

    return Response(_serialize_proposal(proposal))


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def my_department_info(request):
    """
    Return the current user's department info so the frontend can
    pre-select the correct organizer department when creating a proposal.
    Returns: {department_id, department_name, department_code}
    """
    user = request.user
    dept = None
    try:
        sp = user.staff_profile
        dept = (
            sp.get_current_department()
            if hasattr(sp, 'get_current_department')
            else getattr(sp, 'department', None)
        )
    except Exception:
        pass

    if dept:
        return Response({
            'department_id': dept.id,
            'department_name': dept.name,
            'department_code': getattr(dept, 'code', ''),
        })
    return Response({'department_id': None, 'department_name': '', 'department_code': ''})


# ──────────────────────────────────────────────────────────────────────────────
# DETAIL
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def proposal_detail(request, proposal_id):
    try:
        proposal = EventProposal.objects.get(id=proposal_id)
    except EventProposal.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

    original_payload = proposal.proposal_data if isinstance(proposal.proposal_data, dict) else {}
    normalized_payload = _build_doc_payload_for_proposal(proposal, original_payload)
    payload_changed = normalized_payload != original_payload
    if payload_changed:
        proposal.proposal_data = normalized_payload
        proposal.save(update_fields=['proposal_data'])

    needs_doc_regen = payload_changed and (
        not str(original_payload.get('from_name') or '').strip()
        or not any(
            str(original_payload.get(key) or '').strip()
            for key in ('organizer_department_doc', 'organizer_department_raw', 'organizer_department', 'department')
        )
    )
    if needs_doc_regen:
        if proposal.status == EventProposal.Status.HAA_APPROVED and proposal.haa_approved_by_name:
            _regenerate_approved_doc(proposal, request)
        elif proposal.hod_approved_by_name:
            _regenerate_doc_after_hod(proposal, request)
        proposal.refresh_from_db()

    return Response(_serialize_proposal(proposal))


# ──────────────────────────────────────────────────────────────────────────────
# BRANDING → Forward to HOD
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def branding_forward(request, proposal_id):
    """Branding reviews the poster/proposal and forwards to the department HOD."""
    perms = get_user_permissions(request.user)
    if 'events.branding_review' not in perms:
        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

    try:
        proposal = EventProposal.objects.select_related('department').get(id=proposal_id)
    except EventProposal.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

    if proposal.status != EventProposal.Status.FORWARDED_TO_BRANDING:
        return Response({'error': f'Cannot forward from status: {proposal.status}'},
                        status=status.HTTP_400_BAD_REQUEST)

    now = timezone.now()
    user = request.user
    note = request.data.get('note', '')

    proposal.status = EventProposal.Status.FORWARDED_TO_HOD
    proposal.branding_reviewed_by = user
    proposal.branding_reviewed_by_name = _user_display_name(user)
    proposal.branding_reviewed_at = now
    proposal.branding_note = note
    proposal.save()

    # Notify the HOD of that department
    hod_user = _get_hod_for_department(proposal.department)
    if hod_user:
        _create_notification(
            hod_user,
            f'Event Proposal for Approval: {proposal.title}',
            f'Branding has reviewed the event proposal "{proposal.title}" and forwarded it to you.',
            link='/hod/event-approvals',
            data={'proposal_id': str(proposal.id)},
        )

    # Also notify the staff creator
    _create_notification(
        proposal.created_by,
        f'Proposal Forwarded to HOD: {proposal.title}',
        f'Your event proposal "{proposal.title}" has been reviewed by Branding and forwarded to the HOD.',
        data={'proposal_id': str(proposal.id)},
    )

    return Response(_serialize_proposal(proposal))


# ──────────────────────────────────────────────────────────────────────────────
# HOD → Approve → Forward to HAA
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def hod_approve(request, proposal_id):
    """HOD approves the proposal and it auto-forwards to HAA."""
    perms = get_user_permissions(request.user)
    if 'events.hod_approve' not in perms:
        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

    try:
        proposal = EventProposal.objects.select_related('department').get(id=proposal_id)
    except EventProposal.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

    if proposal.status != EventProposal.Status.FORWARDED_TO_HOD:
        return Response({'error': f'Cannot approve from status: {proposal.status}'},
                        status=status.HTTP_400_BAD_REQUEST)

    # HOD must manage the proposal's department. This is a soft check;
    # if the user has no department mappings, we allow them to proceed
    # as they may be a superuser or have cross-departmental duties.
    user = request.user
    try:
        hod_depts = _get_user_hod_departments(user)
        if proposal.department and hod_depts:
            hod_dept_ids = {d.id for d in hod_depts}
            if proposal.department_id not in hod_dept_ids:
                return Response(
                    {'error': "You are not an HOD for this proposal's department"},
                    status=status.HTTP_403_FORBIDDEN,
                )
    except Exception:
        # In case of error fetching departments, allow to proceed
        pass

    now = timezone.now()
    note = request.data.get('note', '')

    proposal.status = EventProposal.Status.FORWARDED_TO_HAA
    proposal.hod_approved_by = user
    proposal.hod_approved_by_name = _user_display_name(user)
    proposal.hod_approved_at = now
    proposal.hod_note = note
    proposal.save()

    # ── Regenerate doc with HOD name so HAA sees it pre-filled ─────────
    _regenerate_doc_after_hod(proposal, request)
    proposal.refresh_from_db()  # pick up new doc URL for notifications

    # Notify HAA
    haa_user = _get_haa_user()
    if haa_user:
        _create_notification(
            haa_user,
            f'Event Proposal for Final Approval: {proposal.title}',
            f'HOD {_user_display_name(user)} has approved the event proposal "{proposal.title}". '
            f'The proposal form now includes HOD approval. Awaiting your final approval.',
            link='/haa/event-approvals',
            data={
                'proposal_id': str(proposal.id),
                'proposal_doc_url': proposal.proposal_doc_url,
                'poster_url': proposal.poster_url,
            },
        )

    # Notify creator
    _create_notification(
        proposal.created_by,
        f'HOD Approved: {proposal.title}',
        f'Your event proposal "{proposal.title}" has been approved by the HOD and forwarded to the HAA.',
        data={'proposal_id': str(proposal.id)},
    )

    return Response(_serialize_proposal(proposal))


# ──────────────────────────────────────────────────────────────────────────────
# HAA → Final Approve
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def haa_approve(request, proposal_id):
    """HAA gives final approval. Triggers notification to the creator with docs."""
    perms = get_user_permissions(request.user)
    if 'events.haa_approve' not in perms:
        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

    try:
        proposal = EventProposal.objects.get(id=proposal_id)
    except EventProposal.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

    if proposal.status != EventProposal.Status.FORWARDED_TO_HAA:
        return Response({'error': f'Cannot approve from status: {proposal.status}'},
                        status=status.HTTP_400_BAD_REQUEST)

    user = request.user
    now = timezone.now()
    note = request.data.get('note', '')

    proposal.status = EventProposal.Status.HAA_APPROVED
    proposal.haa_approved_by = user
    proposal.haa_approved_by_name = _user_display_name(user)
    proposal.haa_approved_at = now
    proposal.haa_note = note
    proposal.save()

    # ── Regenerate the proposal document with both approval names ──────
    _regenerate_approved_doc(proposal, request)
    proposal.refresh_from_db()  # pick up final doc URL

    # ── Notify the creator ──────────────────────────────────────────────
    _create_notification(
        proposal.created_by,
        f'Event Approved! {proposal.title}',
        (
            f'Your event proposal "{proposal.title}" has received final approval from HAA. '
            f'Download your approved documents from the event details page.'
        ),
        link=f'/events/my-proposals',
        data={
            'proposal_id': str(proposal.id),
            'poster_url': proposal.poster_url,
            'proposal_doc_url': proposal.proposal_doc_url,
        },
    )

    # ── Notify branding team with final approved poster ─────────────────
    _notify_role_users(
        'BRANDING',
        f'Event Fully Approved — Final Poster: {proposal.title}',
        f'The event proposal "{proposal.title}" (HOD: {proposal.hod_approved_by_name}, '
        f'HAA: {proposal.haa_approved_by_name}) has been fully approved. '
        f'Please note the final approved poster for your records.',
        link='/branding/list-posters',
        data={
            'proposal_id': str(proposal.id),
            'poster_url': proposal.poster_url,
            'proposal_doc_url': proposal.proposal_doc_url,
        },
    )

    # ── Notify the HOD with the final approved form ─────────────────────
    hod_user = _get_hod_for_department(proposal.department)
    if hod_user:
        _create_notification(
            hod_user,
            f'Event Proposal Fully Approved: {proposal.title}',
            f'The event proposal "{proposal.title}" that you approved has now received '
            f'final approval from HAA ({proposal.haa_approved_by_name}). '
            f'The final approved form is ready.',
            link='/hod/event-approvals',
            data={
                'proposal_id': str(proposal.id),
                'poster_url': proposal.poster_url,
                'proposal_doc_url': proposal.proposal_doc_url,
            },
        )

    # ── Final approval communications (email + WhatsApp) ───────────────
    _send_final_approval_notifications(proposal)

    return Response(_serialize_proposal(proposal))


# ──────────────────────────────────────────────────────────────────────────────
# REJECT (any reviewer can reject)
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def proposal_reject(request, proposal_id):
    """Reject a proposal at any stage. Requires at least one review permission."""
    perms = get_user_permissions(request.user)
    allowed_perms = {'events.branding_review', 'events.hod_approve', 'events.haa_approve'}
    if not perms & allowed_perms:
        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

    try:
        proposal = EventProposal.objects.get(id=proposal_id)
    except EventProposal.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

    if proposal.status in (EventProposal.Status.HAA_APPROVED, EventProposal.Status.REJECTED):
        return Response({'error': 'Cannot reject an already finalized proposal'},
                        status=status.HTTP_400_BAD_REQUEST)

    reason = request.data.get('reason', '').strip()
    if not reason:
        return Response({'error': 'reason is required'}, status=status.HTTP_400_BAD_REQUEST)

    now = timezone.now()
    user = request.user

    proposal.status = EventProposal.Status.REJECTED
    proposal.rejection_reason = reason
    proposal.rejected_by = user
    proposal.rejected_at = now
    proposal.save()

    # Notify the creator
    _create_notification(
        proposal.created_by,
        f'Proposal Rejected: {proposal.title}',
        f'Your event proposal "{proposal.title}" was rejected by {_user_display_name(user)}. Reason: {reason}',
        data={'proposal_id': str(proposal.id)},
    )

    return Response(_serialize_proposal(proposal))


# ──────────────────────────────────────────────────────────────────────────────
# NOTIFICATIONS
# ──────────────────────────────────────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def notifications_list(request):
    """Return the latest 50 notifications for the current user."""
    try:
        from staff_requests.views import run_semester_policy_maintenance
        run_semester_policy_maintenance()
    except Exception:
        pass

    qs = UserNotification.objects.filter(user=request.user).order_by('-created_at')[:50]
    return Response([
        {
            'id': n.id,
            'title': n.title,
            'message': n.message,
            'link': n.link,
            'read': n.read,
            'data': n.data,
            'created_at': n.created_at.isoformat(),
        }
        for n in qs
    ])


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def notification_mark_read(request, notification_id):
    """Mark a notification as read."""
    try:
        n = UserNotification.objects.get(id=notification_id, user=request.user)
    except UserNotification.DoesNotExist:
        return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
    n.read = True
    n.save(update_fields=['read'])
    return Response({'ok': True})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def notifications_unread_count(request):
    """Return the count of unread notifications."""
    try:
        from staff_requests.views import run_semester_policy_maintenance
        run_semester_policy_maintenance()
    except Exception:
        pass

    count = UserNotification.objects.filter(user=request.user, read=False).count()
    return Response({'count': count})


# ──────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ──────────────────────────────────────────────────────────────────────────────

def _notify_role_users(role_name: str, title: str, message: str, link='', data=None):
    """Send an in-app notification to every user with *role_name*."""
    user_ids = (
        UserRole.objects
        .filter(role__name=role_name)
        .values_list('user_id', flat=True)
    )
    notifications = [
        UserNotification(
            user_id=uid,
            title=title,
            message=message,
            link=link,
            data=data or {},
        )
        for uid in user_ids
    ]
    UserNotification.objects.bulk_create(notifications)


def _regenerate_approved_doc(proposal: EventProposal, request=None):
    """
    Re-generate the proposal DOCX with both HOD and HAA approval names filled in.
    Called after HAA gives final approval.
    """
    try:
        from template_api.services.event_proposal_docx import generate_event_proposal_docx

        pd = _build_doc_payload_for_proposal(proposal, proposal.proposal_data)
        pd['hod_approved_by_name'] = proposal.hod_approved_by_name
        pd['haa_approved_by_name'] = proposal.haa_approved_by_name
        pd['hod_approved_at'] = proposal.hod_approved_at
        pd['haa_approved_at'] = proposal.haa_approved_at
        pd['generated_on'] = timezone.now()

        # Use the real request so build_absolute_uri works with the actual host
        if request is None:
            from django.test import RequestFactory
            request = RequestFactory().post('/fake/')
            request.user = proposal.created_by

        result = generate_event_proposal_docx(request, pd)
        if result and hasattr(result, 'get'):
            doc_url = result.get('download_url') or result.get('url', '')
            doc_name = result.get('filename') or result.get('name', '')
            if doc_url:
                proposal.proposal_doc_url = doc_url
            if doc_name:
                proposal.proposal_doc_name = doc_name
            proposal.save(update_fields=['proposal_doc_url', 'proposal_doc_name'])

    except Exception as e:
        logger.warning(f'Failed to regenerate approved doc for proposal {proposal.id}: {e}')


def _regenerate_doc_after_hod(proposal: EventProposal, request=None):
    """
    Re-generate the DOCX with the HOD approval name filled in.
    Called right after HOD approves so HAA sees the form with HOD already signed.
    HAA section is intentionally left blank — it will be filled on final approval.
    """
    try:
        from template_api.services.event_proposal_docx import generate_event_proposal_docx

        pd = _build_doc_payload_for_proposal(proposal, proposal.proposal_data)
        pd['hod_approved_by_name'] = proposal.hod_approved_by_name
        pd['haa_approved_by_name'] = ''   # HAA not yet approved
        pd['hod_approved_at'] = proposal.hod_approved_at
        pd['haa_approved_at'] = ''
        pd['generated_on'] = timezone.now()

        # Use the real request so build_absolute_uri works with the actual host
        if request is None:
            from django.test import RequestFactory
            request = RequestFactory().post('/fake/')
            request.user = proposal.created_by

        result = generate_event_proposal_docx(request, pd)
        if result and hasattr(result, 'get'):
            doc_url = result.get('download_url') or result.get('url', '')
            doc_name = result.get('filename') or result.get('name', '')
            if doc_url:
                proposal.proposal_doc_url = doc_url
            if doc_name:
                proposal.proposal_doc_name = doc_name
            proposal.save(update_fields=['proposal_doc_url', 'proposal_doc_name'])

    except Exception as e:
        logger.warning(f'Failed to regenerate doc after HOD approval for proposal {proposal.id}: {e}')


def _normalize_mobile_number(raw_value: str) -> str:
    digits = ''.join(ch for ch in str(raw_value or '') if ch.isdigit())
    if not digits:
        return ''
    if len(digits) == 10:
        return f'+91{digits}'
    if len(digits) == 12 and digits.startswith('91'):
        return f'+{digits}'
    if str(raw_value).strip().startswith('+'):
        return str(raw_value).strip()
    return f'+{digits}'


def _send_final_approval_email(proposal: EventProposal):
    recipient = str(getattr(proposal.created_by, 'email', '') or '').strip()
    if not recipient:
        logger.info('No email address for user %s, skipping final approval email', proposal.created_by_id)
        return

    subject = f'Event Approved: {proposal.title}'
    message = (
        f'Your event proposal "{proposal.title}" has been fully approved.\n\n'
        f'HOD: {proposal.hod_approved_by_name}\n'
        f'HAA: {proposal.haa_approved_by_name}\n\n'
        f'Poster: {proposal.poster_url}\n'
        f'Proposal Document: {proposal.proposal_doc_url}\n\n'
        f'Please check your IDCS Event Proposals page for full details.'
    )

    from_email = str(getattr(settings, 'DEFAULT_FROM_EMAIL', '') or getattr(settings, 'EMAIL_HOST_USER', '') or '').strip() or None
    send_mail(subject, message, from_email, [recipient], fail_silently=False)
    logger.info('Final approval email sent for proposal %s to %s', proposal.id, recipient)


def _send_whatsapp_notification(proposal: EventProposal):
    """
    Attempt to send WhatsApp / SMS notification for final approval.
    Uses the whatsapp-server if available.
    """
    try:
        import requests as http_requests

        # Get the creator's mobile number
        mobile = ''
        try:
            mobile = proposal.created_by.staff_profile.mobile_number
        except Exception:
            mobile = proposal.created_by.mobile_no

        if not mobile:
            logger.info(f'No mobile number for user {proposal.created_by_id}, skipping WhatsApp')
            return

        normalized_mobile = _normalize_mobile_number(mobile)
        if not normalized_mobile:
            logger.info(f'Invalid mobile number for user {proposal.created_by_id}, skipping WhatsApp')
            return

        whatsapp_url = getattr(settings, 'WHATSAPP_SERVER_URL', 'http://localhost:3001')
        message = (
            f'🎉 Your event proposal "{proposal.title}" has been fully approved!\n\n'
            f'HOD: {proposal.hod_approved_by_name}\n'
            f'HAA: {proposal.haa_approved_by_name}\n\n'
            f'You can download the approved documents from the IDCS portal.'
        )

        http_requests.post(
            f'{whatsapp_url}/api/send-message',
            json={'phone': normalized_mobile, 'message': message},
            timeout=10,
        )
        logger.info(f'WhatsApp notification sent for proposal {proposal.id} to {normalized_mobile}')

    except Exception as e:
        logger.warning(f'WhatsApp notification failed for proposal {proposal.id}: {e}')


def _send_final_approval_notifications(proposal: EventProposal):
    try:
        _send_final_approval_email(proposal)
    except Exception as exc:
        logger.warning('Final approval email failed for proposal %s: %s', proposal.id, exc)

    try:
        _send_whatsapp_notification(proposal)
    except Exception as exc:
        logger.warning('Final approval WhatsApp failed for proposal %s: %s', proposal.id, exc)
