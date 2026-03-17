from __future__ import annotations

import logging
import re
import uuid
from datetime import datetime
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils import timezone

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except Exception:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_BREAK = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    WD_TAB_ALIGNMENT = None
    Inches = None
    Pt = None
    qn = None


logger = logging.getLogger(__name__)

DOCX_FILE_NAME = 'Event Proposal Format.docx'
DOCX_STORAGE_FILE_NAME = 'event-proposal-format.docx'
DOCX_TEMPLATE_CANDIDATES = (
    'Event Proposal Format FInal.docx',
    'Event Proposal Format.docx',
)

IST_TIMEZONE = ZoneInfo('Asia/Kolkata')


def _project_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _template_path() -> Path:
    candidates: list[Path] = []
    for file_name in DOCX_TEMPLATE_CANDIDATES:
        candidates.extend(
            [
                _project_root() / 'frontend' / 'src' / 'assets' / file_name,
                _project_root() / 'backend' / 'staticfiles' / 'assets' / file_name,
            ]
        )

    for candidate in candidates:
        if candidate.exists():
            return candidate

    raise FileNotFoundError(f'Could not find any DOCX template from: {", ".join(DOCX_TEMPLATE_CANDIDATES)}')


def _norm(text: Any) -> str:
    return re.sub(r'\s+', ' ', str(text or '')).strip()


def _format_ist_datetime(value: Any) -> str:
    if value in (None, ''):
        return ''

    parsed: datetime | None = None
    parsed_is_ist_text = False
    if isinstance(value, datetime):
        parsed = value
    else:
        raw_text = _norm(value)
        if not raw_text:
            return ''

        try:
            parsed = datetime.fromisoformat(raw_text.replace('Z', '+00:00'))
        except ValueError:
            parsed = None

        if parsed is None:
            for fmt in ('%d-%m-%Y %I:%M %p IST', '%d-%m-%Y %I:%M %p', '%d-%m-%Y', '%d/%m/%Y'):
                try:
                    parsed = datetime.strptime(raw_text, fmt)
                    if fmt.endswith('IST'):
                        parsed_is_ist_text = True
                    break
                except ValueError:
                    continue

    if parsed is None:
        return _norm(value)

    if timezone.is_naive(parsed):
        if parsed_is_ist_text:
            parsed = parsed.replace(tzinfo=IST_TIMEZONE)
        else:
            parsed = timezone.make_aware(parsed, timezone.get_current_timezone())
    parsed = timezone.localtime(parsed, IST_TIMEZONE)
    return parsed.strftime('%d-%m-%Y %I:%M %p IST')


def _is_name_placeholder(text: str) -> bool:
    normalized = _norm(text)
    return bool(re.fullmatch(r'(Dr\.?\s*)?[.…•·_]+', normalized))


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(text)


def _set_runs_bold(paragraph, is_bold: bool) -> None:
    try:
        for run in paragraph.runs:
            run.bold = is_bold
    except Exception:
        pass


def _set_tab_stops(paragraph) -> None:
    if Inches is None:
        return
    try:
        p_pr = paragraph._element.get_or_add_pPr()
        existing_tabs = p_pr.find(qn('w:tabs')) if qn is not None else None
        if existing_tabs is not None:
            p_pr.remove(existing_tabs)
        stops = paragraph.paragraph_format.tab_stops
        for position in (Inches(1.55), Inches(4.7), Inches(6.1)):
            stops.add_tab_stop(position)
    except Exception:
        pass


def _set_event_detail_tab_stops(paragraph, two_column: bool = False) -> None:
    if Inches is None:
        return
    try:
        p_pr = paragraph._element.get_or_add_pPr()
        existing_tabs = p_pr.find(qn('w:tabs')) if qn is not None else None
        if existing_tabs is not None:
            p_pr.remove(existing_tabs)
        stops = paragraph.paragraph_format.tab_stops
        stops.add_tab_stop(Inches(2.3))
        stops.add_tab_stop(Inches(2.45))
        if two_column:
            stops.add_tab_stop(Inches(5.65))
            stops.add_tab_stop(Inches(6.95))
    except Exception:
        pass


def _set_zero_spacing(paragraph) -> None:
    if Pt is None:
        return
    try:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    except Exception:
        pass


def _set_page_flow(paragraph, *, keep_with_next: bool = False, keep_together: bool = False, page_break_before: bool = False) -> None:
    try:
        paragraph.paragraph_format.keep_with_next = keep_with_next
        paragraph.paragraph_format.keep_together = keep_together
        paragraph.paragraph_format.page_break_before = page_break_before
    except Exception:
        pass


def _set_line_spacing_single(paragraph) -> None:
    try:
        paragraph.paragraph_format.line_spacing = 1
    except Exception:
        pass


def _set_paragraph_indents(paragraph, *, left=None, right=None, first_line=None) -> None:
    try:
        if left is not None:
            paragraph.paragraph_format.left_indent = left
        if right is not None:
            paragraph.paragraph_format.right_indent = right
        if first_line is not None:
            paragraph.paragraph_format.first_line_indent = first_line
    except Exception:
        pass


def _set_paragraph_alignment(paragraph, alignment: int) -> None:
    try:
        paragraph.alignment = alignment
    except Exception:
        pass


def _set_cell_text(cell, text: str, alignment: int = 0) -> None:
    cell.text = text
    try:
        if WD_CELL_VERTICAL_ALIGNMENT is not None:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass

    for paragraph in cell.paragraphs:
        _set_paragraph_alignment(paragraph, alignment)
        _set_zero_spacing(paragraph)


def _set_signature_heading(
    paragraph,
    hod_name: str = '',
    hod_approved_at: str = '',
    coordinator_lines: list[str] | None = None,
) -> None:
    left_lines = ['Coordinators']
    if coordinator_lines:
        left_lines.extend([_norm(line).upper() for line in coordinator_lines if _norm(line)])

    right_lines: list[str] = []
    if hod_name:
        right_lines = ['HOD', hod_name, 'Approved']
        if hod_approved_at:
            right_lines.append(hod_approved_at)
    elif coordinator_lines:
        right_lines = ['HOD']

    row_count = max(len(left_lines), len(right_lines) or 1)
    merged_lines: list[str] = []
    for idx in range(row_count):
        left = left_lines[idx] if idx < len(left_lines) else ''
        right = right_lines[idx] if idx < len(right_lines) else ''
        merged_lines.append(f'{left}\t{right}')

    _set_paragraph_text(paragraph, '\n'.join(merged_lines))
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    _set_runs_bold(paragraph, False)
    if Inches is None:
        return
    try:
        stops = paragraph.paragraph_format.tab_stops
        if WD_TAB_ALIGNMENT is not None:
            stops.add_tab_stop(Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
        else:
            stops.add_tab_stop(Inches(6.5))
    except Exception:
        pass


def _set_signature_names(paragraph, lines: list[str]) -> None:
    normalized_lines = [_norm(line).upper() for line in lines if _norm(line)]
    text = '\n'.join(normalized_lines)
    _set_paragraph_text(paragraph, text)
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    _set_runs_bold(paragraph, False)
    _set_page_flow(paragraph, keep_with_next=False, keep_together=True, page_break_before=False)


def _set_signoff_text(paragraph, text: str) -> None:
    _set_paragraph_text(paragraph, text)
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    _set_page_flow(paragraph, keep_with_next=False, keep_together=False, page_break_before=False)
    if Inches is not None:
        _set_paragraph_indents(paragraph, left=Inches(4.6), right=Inches(0.1), first_line=Inches(0))


def _set_date_place_text(paragraph, text: str) -> None:
    _set_paragraph_text(paragraph, text)
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    _set_page_flow(paragraph, keep_with_next=False, keep_together=False, page_break_before=False)
    if Inches is not None:
        _set_paragraph_indents(paragraph, left=Inches(0), right=Inches(0), first_line=Inches(0))


def _set_authority_line(paragraph, haa_name: str = '', hod_name: str = '', haa_approved_at: str = '') -> None:
    # Keep all three designations on one fixed line and reserve a blank row
    # above them for signatures (especially IQAC/Principal).
    lines: list[str] = []
    # Keep one controlled spacer row above designations.
    lines.append('\t\t')
    lines.append('HAA\tIQAC\tPRINCIPAL')

    _set_paragraph_text(paragraph, '\n'.join(lines))
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    if Inches is not None:
        _set_paragraph_indents(paragraph, left=Inches(0), right=Inches(0), first_line=Inches(0))
    _set_runs_bold(paragraph, False)
    _set_page_flow(paragraph, keep_with_next=False, keep_together=True, page_break_before=False)
    if Inches is None:
        return
    try:
        # Clear any existing tabs from template to avoid column drift.
        p_pr = paragraph._element.get_or_add_pPr()
        if qn is not None:
            existing_tabs = p_pr.find(qn('w:tabs'))
            if existing_tabs is not None:
                p_pr.remove(existing_tabs)

        stops = paragraph.paragraph_format.tab_stops
        # Use fixed left-aligned tab columns to keep all three designations on
        # the exact same baseline across DOCX renderers.
        # Tuned for LibreOffice rendering so IQAC and PRINCIPAL sit on the
        # same visual line with PRINCIPAL shifted left from the page edge.
        stops.add_tab_stop(Inches(3.00))
        stops.add_tab_stop(Inches(4.50))
    except Exception:
        pass


def _set_haa_approval_block(paragraph, haa_name: str = '', haa_approved_at: str = '') -> None:
    """Render HAA approval details under the authority row, left column only."""
    lines: list[str] = []
    haa_name_text = _norm(haa_name)
    haa_time_text = _norm(haa_approved_at)
    if haa_name_text:
        lines.append(haa_name_text)
        lines.append('Approved')
    if haa_time_text:
        lines.append(haa_time_text)

    _set_paragraph_text(paragraph, '\n'.join(lines) if lines else '')
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    if Inches is not None:
        _set_paragraph_indents(paragraph, left=Inches(0), right=Inches(0), first_line=Inches(0))
    _set_runs_bold(paragraph, False)
    _set_page_flow(paragraph, keep_with_next=False, keep_together=False, page_break_before=False)


def _shrink_blank_paragraph(paragraph) -> None:
    if _norm(paragraph.text):
        return
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)
    if Pt is not None:
        for run in paragraph.runs:
            try:
                run.font.size = Pt(1)
            except Exception:
                pass


def _normalize_first_page_block(document) -> None:
    for idx in (27, 28, 29, 30, 31, 33, 35):
        if idx >= len(document.paragraphs):
            continue
        paragraph = document.paragraphs[idx]
        _set_zero_spacing(paragraph)
        _set_line_spacing_single(paragraph)
        _set_page_flow(paragraph, keep_with_next=False, keep_together=False, page_break_before=False)

    if 32 < len(document.paragraphs):
        budget_heading = document.paragraphs[32]
        _set_zero_spacing(budget_heading)
        _set_line_spacing_single(budget_heading)
        _set_page_flow(budget_heading, keep_with_next=False, keep_together=False, page_break_before=True)

    for idx in (34, 36):
        if idx < len(document.paragraphs):
            _shrink_blank_paragraph(document.paragraphs[idx])

    if 27 < len(document.paragraphs):
        _set_paragraph_alignment(document.paragraphs[27], 1 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.CENTER)
    if 32 < len(document.paragraphs):
        _set_paragraph_alignment(document.paragraphs[32], 1 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.CENTER)
    if 33 < len(document.paragraphs):
        _set_paragraph_alignment(document.paragraphs[33], 1 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.CENTER)
    if 35 < len(document.paragraphs):
        _set_paragraph_alignment(document.paragraphs[35], 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
        if Inches is not None:
            _set_paragraph_indents(document.paragraphs[35], left=Inches(0), right=Inches(0), first_line=Inches(0))


def _set_label_value_line(paragraph, label: str, value: str) -> None:
    _set_event_detail_tab_stops(paragraph)
    _set_paragraph_text(paragraph, f'{label}\t: {value}')
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)


def _set_two_column_line(paragraph, left_label: str, left_value: str, right_label: str, right_value: str) -> None:
    _set_event_detail_tab_stops(paragraph, two_column=True)
    _set_paragraph_text(paragraph, f'{left_label}\t: {left_value}\t{right_label} : {right_value}')
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)


def _set_stacked_label_values(paragraph, pairs: list[tuple[str, str]]) -> None:
    lines = []
    for label, value in pairs:
        if label or value:
            lines.append(f'{label}\t: {value}'.rstrip())
    _set_event_detail_tab_stops(paragraph)
    _set_paragraph_text(paragraph, '\n'.join(lines))
    _set_paragraph_alignment(paragraph, 0 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT)
    _set_zero_spacing(paragraph)
    _set_line_spacing_single(paragraph)


def _set_centered_lines(paragraph, lines: list[str]) -> None:
    text = '\n'.join(line for line in lines if _norm(line))
    _set_paragraph_text(paragraph, text)
    _set_zero_spacing(paragraph)
    try:
        paragraph.alignment = 1
    except Exception:
        pass


def _parse_decimal(value: Any) -> Decimal:
    text = _norm(value).replace(',', '')
    if not text:
        return Decimal('0')
    try:
        return Decimal(text)
    except (InvalidOperation, ValueError):
        return Decimal('0')


def _number_to_words(number: int) -> str:
    ones = {
        0: 'zero', 1: 'one', 2: 'two', 3: 'three', 4: 'four', 5: 'five', 6: 'six', 7: 'seven', 8: 'eight', 9: 'nine',
        10: 'ten', 11: 'eleven', 12: 'twelve', 13: 'thirteen', 14: 'fourteen', 15: 'fifteen', 16: 'sixteen', 17: 'seventeen', 18: 'eighteen', 19: 'nineteen',
    }
    tens = {
        20: 'twenty', 30: 'thirty', 40: 'forty', 50: 'fifty', 60: 'sixty', 70: 'seventy', 80: 'eighty', 90: 'ninety',
    }

    if number < 20:
        return ones[number]
    if number < 100:
        return tens[number - (number % 10)] + (f' {ones[number % 10]}' if number % 10 else '')
    if number < 1000:
        return ones[number // 100] + ' hundred' + (f' {_number_to_words(number % 100)}' if number % 100 else '')

    scales = [
        (1_000_000_000, 'billion'),
        (1_000_000, 'million'),
        (1_000, 'thousand'),
    ]
    for divisor, label in scales:
        if number >= divisor:
            major, remainder = divmod(number, divisor)
            return _number_to_words(major) + f' {label}' + (f' {_number_to_words(remainder)}' if remainder else '')

    return str(number)


def _amount_to_words(value: Any) -> str:
    amount = _parse_decimal(value)
    whole = int(amount)
    fraction = int((amount - whole) * 100)
    words = _number_to_words(whole)
    if fraction > 0:
        return f'Rupees: {words} and {_number_to_words(fraction)} paise'
    return f'Rupees: {words}'


def _clean_department_name(value: str) -> str:
    text = _norm(value)
    if not text:
        return ''
    prefix = 'DEPARTMENT OF '
    if text.upper().startswith(prefix):
        text = text[len(prefix):].strip()
    return text


def _build_date_range(data: dict[str, Any]) -> tuple[str, str]:
    from_date = _norm(data.get('from_date') or data.get('event_date_from'))
    to_date = _norm(data.get('to_date') or data.get('event_date_to'))
    if from_date or to_date:
        return from_date, to_date

    start_day = _norm(data.get('start_day'))
    end_day = _norm(data.get('end_day'))
    start_month = _norm(data.get('start_month'))
    year = _norm(data.get('year'))

    if start_day or start_month or year:
        left = ' '.join(part for part in [start_day, start_month, year] if part)
        right = ' '.join(part for part in [end_day or start_day, start_month, year] if part)
        return left, right

    event_date = _norm(data.get('event_date'))
    return event_date, event_date


def _build_affiliation(data: dict[str, Any]) -> str:
    explicit = _norm(data.get('resource_person_affiliation') or data.get('affiliation'))
    if explicit:
        return explicit
    parts = [
        _norm(data.get('chief_guest_company')),
        _norm(data.get('chief_guest_location')),
    ]
    return ', '.join(part for part in parts if part)


def _split_person_values(value: Any) -> list[str]:
    text = _norm(value)
    if not text:
        return []
    parts = [part.strip() for part in re.split(r'[\n;,]+', text) if part.strip()]
    return parts or [text]


def _unique_values(values: list[str]) -> list[str]:
    seen: set[str] = set()
    unique: list[str] = []
    for value in values:
        normalized = _norm(value)
        if not normalized:
            continue
        key = normalized.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(normalized)
    return unique


def _normalized_key(key: str) -> str:
    return re.sub(r'[^a-z0-9]+', '_', str(key or '').lower()).strip('_')


def _extract_coordinator_lines(proposal_data: dict[str, Any]) -> list[str]:
    values: list[str] = []

    for key in (
        'coordinator',
        'co_coordinator',
        'coordinator_name',
        'co_coordinator_name',
        'faculty_coordinator_1',
        'faculty_coordinator_2',
        'student_coordinator',
        'coordinators',
    ):
        values.extend(_split_person_values(proposal_data.get(key)))

    for index in range(1, 21):
        values.extend(_split_person_values(proposal_data.get(f'committee_member_{index}_name')))

    for raw_key, raw_value in proposal_data.items():
        key = _normalized_key(raw_key)
        if not _norm(raw_value):
            continue

        if re.fullmatch(r'committee_member_\d+_name', key):
            values.extend(_split_person_values(raw_value))
            continue

        if 'coordinator' not in key:
            continue
        if any(token in key for token in ('role', 'designation', 'position', 'email', 'phone', 'department')):
            continue
        values.extend(_split_person_values(raw_value))

    return _unique_values(values)


def _extract_resource_person_lines(proposal_data: dict[str, Any]) -> list[str]:
    values: list[str] = []

    for key in ('resource_person', 'resource_persons', 'chief_guest_name', 'guest_name'):
        values.extend(_split_person_values(proposal_data.get(key)))

    for index in range(1, 21):
        values.extend(_split_person_values(proposal_data.get(f'chief_guest_{index}_name')))
        values.extend(_split_person_values(proposal_data.get(f'resource_person_{index}')))
        values.extend(_split_person_values(proposal_data.get(f'resource_person_{index}_name')))
        values.extend(_split_person_values(proposal_data.get(f'guest_{index}_name')))

    for raw_key, raw_value in proposal_data.items():
        key = _normalized_key(raw_key)
        if not _norm(raw_value):
            continue

        is_chief_guest_name = key.startswith('chief_guest') and 'name' in key and 'photo' not in key
        is_resource_person_name = (
            key.startswith('resource_person')
            and all(token not in key for token in ('designation', 'position', 'affiliation', 'company', 'location', 'email', 'phone'))
        )
        if is_chief_guest_name or is_resource_person_name:
            values.extend(_split_person_values(raw_value))

    return _unique_values(values)


def _extract_designation_lines(proposal_data: dict[str, Any]) -> list[str]:
    values: list[str] = []

    for key in ('designation', 'chief_guest_position'):
        values.extend(_split_person_values(proposal_data.get(key)))

    for index in range(1, 21):
        values.extend(_split_person_values(proposal_data.get(f'chief_guest_{index}_position')))
        values.extend(_split_person_values(proposal_data.get(f'resource_person_{index}_designation')))
        values.extend(_split_person_values(proposal_data.get(f'guest_{index}_designation')))

    for raw_key, raw_value in proposal_data.items():
        key = _normalized_key(raw_key)
        if not _norm(raw_value):
            continue
        if ('chief_guest' in key or 'resource_person' in key or 'guest' in key) and ('designation' in key or 'position' in key):
            values.extend(_split_person_values(raw_value))

    return _unique_values(values)


def _extract_affiliation_text(proposal_data: dict[str, Any]) -> str:
    values: list[str] = []

    for key in ('resource_person_affiliation', 'affiliation', 'chief_guest_company', 'chief_guest_location'):
        value = _norm(proposal_data.get(key))
        if value:
            values.append(value)

    for index in range(1, 21):
        for key in (
            f'chief_guest_{index}_company',
            f'chief_guest_{index}_location',
            f'chief_guest_{index}_affiliation',
            f'resource_person_{index}_affiliation',
            f'guest_{index}_affiliation',
        ):
            value = _norm(proposal_data.get(key))
            if value:
                values.append(value)

    fallback = _build_affiliation(proposal_data)
    if fallback:
        values.append(fallback)

    unique_values = _unique_values(values)
    return '; '.join(unique_values)


def _cell_text(value: Any) -> str:
    return _norm(value)


def _has_budget_line_values(line: dict[str, Any] | None) -> bool:
    if not isinstance(line, dict):
        return False
    keys = ('category', 'subType', 'qty', 'total', 'notes')
    return any(_cell_text(line.get(key)) for key in keys)


def _has_income_line_values(line: dict[str, Any] | None) -> bool:
    if not isinstance(line, dict):
        return False
    keys = ('source', 'unitPrice', 'qty', 'total')
    return any(_cell_text(line.get(key)) for key in keys)


def _delete_row(table, row_index: int) -> None:
    try:
        row = table.rows[row_index]
        table._tbl.remove(row._tr)
    except Exception:
        pass


def _fill_budget_table(table, proposal_data: dict[str, Any], department_name: str, event_type: str, from_date: str, to_date: str) -> None:
    budget = proposal_data.get('budget') or {}
    income = proposal_data.get('income') or []
    office = proposal_data.get('office') or {}

    expense_defs = [
        ('chief_guest_honorarium', 'Chief Guest Honorarium'),
        ('travelling_allowance', 'Travelling Allowance'),
        ('refreshment_chief_guest', 'Refreshment (Chief Guest)'),
        ('refreshment_external_participants', 'Refreshment (External Participants)'),
        ('lunch_chief_guest', 'Lunch (Chief Guest)'),
        ('lunch_external_participants', 'Lunch (External Participants)'),
        ('prize_award', 'Prize/Award to Participants (per event)'),
        ('local_travel', 'Local Travel'),
        ('decorations', 'Decorations (Only College level events)'),
        ('welcome_kit', 'Welcome Kit'),
        ('certificates', 'Certificates'),
        ('miscellaneous', 'Miscellaneous'),
    ]

    expense_lines: list[tuple[str, dict[str, Any], str]] = []
    for key, default_label in expense_defs:
        line = budget.get(key) or {}
        if not _has_budget_line_values(line):
            continue
        category = _cell_text(line.get('category'))
        label = default_label
        if key == 'chief_guest_honorarium' and category:
            label = f'{default_label} ({category})'
        elif key == 'travelling_allowance' and category:
            label = f'{default_label} ({category})'
        elif key == 'miscellaneous' and _cell_text(line.get('notes')):
            label = f'{default_label} — {_cell_text(line.get("notes"))}'
        expense_lines.append((label, line, key))

    income_lines = [line for line in income if _has_income_line_values(line)]

    total_budget = _cell_text(proposal_data.get('total_budget_amount'))
    total_income = _cell_text(proposal_data.get('total_income_amount'))
    amount_requested = _cell_text(proposal_data.get('amount_requested_from_institute'))

    # Remove all original data rows between the header and totals sections.
    while len(table.rows) > 1 and _cell_text(table.rows[1].cells[1].text) != 'A, Total Budget Amount (₹)':
        _delete_row(table, 1)

    total_budget_row_index = 1
    total_budget_row = table.rows[total_budget_row_index]

    # Insert compacted expense rows before the total budget row.
    for serial_no, (label, line, key) in enumerate(expense_lines, start=1):
        new_row = table.add_row()
        table._tbl.remove(new_row._tr)
        table._tbl.insert(total_budget_row._tr.getparent().index(total_budget_row._tr), new_row._tr)
        _set_cell_text(new_row.cells[0], str(serial_no))
        _set_cell_text(new_row.cells[1], label)
        _set_cell_text(new_row.cells[2], _cell_text(line.get('unitPrice')))
        _set_cell_text(new_row.cells[3], _cell_text(line.get('qty')))
        _set_cell_text(new_row.cells[4], _cell_text(line.get('total')))

    total_budget_row = table.rows[len(expense_lines) + 1]
    _set_cell_text(total_budget_row.cells[0], str(len(expense_lines) + 1))
    _set_cell_text(total_budget_row.cells[1], 'A, Total Budget Amount (₹)')
    if len(total_budget_row.cells) > 2 and total_budget_row.cells[2]._tc is not total_budget_row.cells[1]._tc:
        _set_cell_text(total_budget_row.cells[2], '')
    if len(total_budget_row.cells) > 3 and total_budget_row.cells[3]._tc is not total_budget_row.cells[1]._tc:
        _set_cell_text(total_budget_row.cells[3], '')
    _set_cell_text(total_budget_row.cells[4], total_budget)

    # Normalize income section row numbering and contents.
    income_header_index = len(expense_lines) + 2
    income_data_start = income_header_index + 1

    while len(table.rows) > income_data_start and _cell_text(table.rows[income_data_start].cells[1].text) != 'B. Total Income':
        _delete_row(table, income_data_start)

    total_income_row = table.rows[income_data_start]
    for serial_no, line in enumerate(income_lines, start=1):
        new_row = table.add_row()
        table._tbl.remove(new_row._tr)
        table._tbl.insert(total_income_row._tr.getparent().index(total_income_row._tr), new_row._tr)
        _set_cell_text(new_row.cells[0], str(serial_no))
        _set_cell_text(new_row.cells[1], _cell_text(line.get('source')))
        _set_cell_text(new_row.cells[2], _cell_text(line.get('unitPrice')))
        _set_cell_text(new_row.cells[3], _cell_text(line.get('qty')))
        _set_cell_text(new_row.cells[4], _cell_text(line.get('total')))

    total_income_row = table.rows[income_data_start + len(income_lines)]
    _set_cell_text(total_income_row.cells[0], '')
    _set_cell_text(total_income_row.cells[1], 'B. Total Income')
    _set_cell_text(total_income_row.cells[2], '')
    _set_cell_text(total_income_row.cells[3], '')
    _set_cell_text(total_income_row.cells[4], total_income)

    amount_row_index = income_data_start + len(income_lines) + 1
    if len(table.rows) > amount_row_index:
        amount_row = table.rows[amount_row_index]
        _set_cell_text(amount_row.cells[0], '')
        _set_cell_text(amount_row.cells[1], 'Amount Requested from the Institute')
        _set_cell_text(amount_row.cells[2], '')
        _set_cell_text(amount_row.cells[3], '')
        _set_cell_text(amount_row.cells[4], amount_requested)


def _fill_office_use_table(table, proposal_data: dict[str, Any], department_name: str, event_type: str, from_date: str, to_date: str) -> None:
    office = proposal_data.get('office') or {}
    budget = proposal_data.get('budget') or {}
    honorarium_category = _cell_text((budget.get('chief_guest_honorarium') or {}).get('category'))
    approved = _cell_text(proposal_data.get('amount_requested_from_institute')) or _cell_text(proposal_data.get('total_budget_amount'))
    if from_date or to_date:
        event_date_text = ' to '.join(part for part in [from_date, to_date] if part)
    else:
        event_date_text = _cell_text(proposal_data.get('event_date'))

    column_count = len(table.rows[0].cells) if table.rows else 0

    if column_count >= 6:
        if len(table.rows) > 0:
            _set_cell_text(table.rows[0].cells[2], department_name)
            _set_cell_text(table.rows[0].cells[4], event_type)
        if len(table.rows) > 1:
            _set_cell_text(table.rows[1].cells[2], event_date_text)
            _set_cell_text(table.rows[1].cells[4], honorarium_category or _cell_text(office.get('expertCategory')))
        if len(table.rows) > 2:
            _set_cell_text(table.rows[2].cells[2], _cell_text(office.get('modeOfEvent')))
            _set_cell_text(table.rows[2].cells[4], _cell_text(office.get('isEventRepeated')))
        if len(table.rows) > 3:
            _set_cell_text(table.rows[3].cells[1], '')
        if len(table.rows) > 4:
            _set_cell_text(table.rows[4].cells[5], approved)
        return

    if len(table.rows) > 0:
        _set_cell_text(table.rows[0].cells[1], department_name)
        _set_cell_text(table.rows[0].cells[3], event_type)
    if len(table.rows) > 1:
        _set_cell_text(table.rows[1].cells[1], event_date_text)
        _set_cell_text(table.rows[1].cells[3], honorarium_category or _cell_text(office.get('expertCategory')))
    if len(table.rows) > 2:
        _set_cell_text(table.rows[2].cells[1], _cell_text(office.get('modeOfEvent')))
        _set_cell_text(table.rows[2].cells[3], _cell_text(office.get('isEventRepeated')))
    if len(table.rows) > 3:
        _set_cell_text(table.rows[3].cells[1], _cell_text(office.get('addressedPOs')))
        _set_cell_text(table.rows[3].cells[3], _cell_text(office.get('copoAttainment')))
    if len(table.rows) > 4:
        for cell in table.rows[4].cells:
            if cell.paragraphs:
                cell.paragraphs[0].text = 'Remarks'
                _set_paragraph_alignment(cell.paragraphs[0], 1 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.CENTER)
                _set_zero_spacing(cell.paragraphs[0])
                for paragraph in cell.paragraphs[1:]:
                    paragraph.text = ''
                    _set_zero_spacing(paragraph)
    if len(table.rows) > 5:
        _set_cell_text(table.rows[5].cells[0], f'Approved Budget\n{approved}', alignment=1 if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.CENTER)


def _fill_addressed_po_table(table, proposal_data: dict[str, Any]) -> None:
    office = proposal_data.get('office') or {}

    if len(table.rows) > 0 and len(table.rows[0].cells) >= 4:
        _set_cell_text(table.rows[0].cells[0], _cell_text(table.rows[0].cells[0].text))
        _set_cell_text(table.rows[0].cells[1], _cell_text(office.get('addressedPOs')))
        _set_cell_text(table.rows[0].cells[2], _cell_text(table.rows[0].cells[2].text))
        _set_cell_text(table.rows[0].cells[3], _cell_text(office.get('copoAttainment')))


def generate_event_proposal_docx(request, proposal_data: dict[str, Any]) -> dict[str, str]:
    if Document is None:
        raise RuntimeError('python-docx is not installed')

    template_path = _template_path()
    document = Document(str(template_path))

    from_name = _norm(proposal_data.get('from_name') or proposal_data.get('from'))
    department_name = _clean_department_name(
        proposal_data.get('organizer_department_doc')
        or proposal_data.get('organizer_department_raw')
        or proposal_data.get('organizer_department')
        or proposal_data.get('department')
    )
    event_type = _norm(proposal_data.get('event_type') or proposal_data.get('nature'))
    event_title = _norm(proposal_data.get('event_title') or proposal_data.get('event_name') or proposal_data.get('title'))
    participants = _norm(proposal_data.get('participants'))
    coordinator_lines = _extract_coordinator_lines(proposal_data)
    if not coordinator_lines:
        coordinator_lines = ['-']

    resource_person_lines = _extract_resource_person_lines(proposal_data)
    designation_lines = _extract_designation_lines(proposal_data)
    affiliation = _extract_affiliation_text(proposal_data)

    if not resource_person_lines:
        resource_person = '-'
        designation = '-'
        affiliation = '-'
    else:
        resource_person = '; '.join(resource_person_lines)
        designation = '; '.join(designation_lines) if designation_lines else '-'
        affiliation = affiliation or '-'
    generated_on = _format_ist_datetime(proposal_data.get('generated_on')) or _format_ist_datetime(timezone.now())
    from_date, to_date = _build_date_range(proposal_data)
    amount_requested = _cell_text(proposal_data.get('amount_requested_from_institute'))

    # Approval names (filled when re-generating after final approval)
    haa_approved_by_name = _norm(proposal_data.get('haa_approved_by_name', ''))
    hod_approved_by_name = _norm(proposal_data.get('hod_approved_by_name', ''))
    haa_approved_at = _format_ist_datetime(proposal_data.get('haa_approved_at'))
    hod_approved_at = _format_ist_datetime(proposal_data.get('hod_approved_at'))

    coordinator_heading_idx: int | None = None
    office_use_idx: int | None = None
    authority_idx: int | None = None
    authority_details_idx: int | None = None
    guidelines_idx: int | None = None

    for idx, paragraph in enumerate(document.paragraphs):
        text = _norm(paragraph.text)
        if not text:
            continue

        if _is_name_placeholder(text) and from_name:
            _set_paragraph_text(paragraph, from_name)
        elif text == '__________________ Engineering,' and department_name:
            _set_paragraph_text(paragraph, f'{department_name} Engineering,')
        elif text.startswith('The department of ____________________ engineering is planning to organize') and department_name:
            _set_paragraph_text(
                paragraph,
                f'The department of {department_name} engineering is planning to organize the following event in our department/college. So, I kindly request you to approve the budget (as attached).',
            )
        elif text.startswith('Nature') and event_type:
            _set_label_value_line(paragraph, 'Nature', event_type)
        elif text.startswith('Title of Event') and event_title:
            _set_label_value_line(paragraph, 'Title of Event', event_title)
        elif text.startswith('Event Date') and (from_date or to_date):
            _set_label_value_line(paragraph, 'Event Date', f'From {from_date} to {to_date}')
        elif text.startswith('Participants') and participants:
            _set_label_value_line(paragraph, 'Participants', participants)
        elif text.startswith('Co-ordinator'):
            coordinator_pairs: list[tuple[str, str]] = []
            for line_index, line_value in enumerate(coordinator_lines):
                if line_index == 0:
                    label = 'Co-ordinator'
                elif line_index == 1:
                    label = 'Co-Coordinator'
                else:
                    label = f'Co-Coordinator {line_index}'
                coordinator_pairs.append((label, line_value))
            _set_stacked_label_values(paragraph, coordinator_pairs)
        elif text.startswith('Resource Person(s):'):
            _set_two_column_line(paragraph, 'Resource Person', resource_person, 'Designation', designation)
        elif text.startswith('Affiliation of the Resource Person:'):
            _set_label_value_line(paragraph, 'Affiliation of the Resource Person', affiliation)
        elif text.startswith('Coordinators'):
            coordinator_heading_idx = idx
            _set_signature_heading(
                paragraph,
                hod_name=hod_approved_by_name,
                hod_approved_at=hod_approved_at,
                coordinator_lines=coordinator_lines,
            )
        elif text == '(for office use only)':
            office_use_idx = idx
            # Shrink any placeholders between signature block and office-use
            # marker so they do not introduce visual gap.
            if coordinator_heading_idx is not None and idx > coordinator_heading_idx + 1:
                for blank_idx in range(coordinator_heading_idx + 1, idx):
                    _shrink_blank_paragraph(document.paragraphs[blank_idx])
        elif (
            text == 'HAA IQAC Principal'
            or (
                text.startswith('HAA')
                and 'iqac' in text.lower()
                and 'principal' in text.lower()
            )
        ):
            authority_idx = idx
            _set_authority_line(
                paragraph,
                haa_name=haa_approved_by_name,
                hod_name=hod_approved_by_name,
                haa_approved_at=haa_approved_at,
            )
            # Put approval details in the next paragraph so the designation row
            # stays fixed on one baseline.
            if idx + 1 < len(document.paragraphs):
                authority_details_idx = idx + 1
                _set_haa_approval_block(
                    document.paragraphs[authority_details_idx],
                    haa_name=haa_approved_by_name,
                    haa_approved_at=haa_approved_at,
                )
        elif text == 'GUIDELINES':
            guidelines_idx = idx
            paragraph.paragraph_format.page_break_before = True
        elif text == '(Dr………………………….)' and from_name:
            _set_signoff_text(paragraph, from_name)
        elif text == 'Yours Sincerely,':
            _set_signoff_text(paragraph, 'Yours Sincerely,')
        elif text == 'Date:':
            _set_date_place_text(paragraph, f'Date: {generated_on}')
        elif text == 'Place: Tiruchirappalli':
            _set_date_place_text(paragraph, 'Place: Tiruchirappalli')
        elif text.startswith('Department of __') and department_name:
            _set_paragraph_text(paragraph, f'Department of {department_name}')
        elif text.startswith('(Rupees'):
            _set_paragraph_text(paragraph, _amount_to_words(amount_requested))

    if office_use_idx is not None and guidelines_idx is not None:
        for shrink_idx in range(office_use_idx + 1, guidelines_idx):
            paragraph = document.paragraphs[shrink_idx]
            if shrink_idx in {authority_idx, authority_details_idx}:
                _set_zero_spacing(paragraph)
            else:
                _shrink_blank_paragraph(paragraph)

    _normalize_first_page_block(document)

    try:
        if len(document.tables) > 0:
            _fill_budget_table(document.tables[0], proposal_data, department_name, event_type, from_date, to_date)
        if len(document.tables) > 2:
            _fill_addressed_po_table(document.tables[1], proposal_data)
            _fill_office_use_table(document.tables[2], proposal_data, department_name, event_type, from_date, to_date)
        elif len(document.tables) > 1:
            _fill_office_use_table(document.tables[1], proposal_data, department_name, event_type, from_date, to_date)
    except Exception as exc:
        logger.warning('Could not fill proposal DOCX tables: %s', exc)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    storage_path = default_storage.save(
        f'proposal-docs/{uuid.uuid4().hex}/{DOCX_STORAGE_FILE_NAME}',
        ContentFile(buffer.getvalue()),
    )
    doc_id = storage_path.split('/')[1] if '/' in storage_path else ''
    abs_url = request.build_absolute_uri(f'/api/canva/proposal-docs/{doc_id}/{DOCX_STORAGE_FILE_NAME}')

    return {
        'name': DOCX_FILE_NAME,
        'filename': DOCX_FILE_NAME,
        'path': storage_path,
        'doc_id': doc_id,
        'url': abs_url,
        'download_url': abs_url,
    }