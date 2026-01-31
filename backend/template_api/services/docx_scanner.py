import base64
import re
import logging
import mimetypes
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from django.utils.text import slugify

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


_NS = {
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


RE_NUMBERED = re.compile(r'^\s*\d+[).]?\s*(.+)$')
RE_Q_PREFIX = re.compile(r'^\s*(?:Q\.|Q:|Question\s*\d*[).]?)\s*(.+)$', re.IGNORECASE)
RE_OPT_1 = re.compile(r'^\s*[A-Da-d][.)]\s+(.+)$')
RE_OPT_2 = re.compile(r'^\s*\([a-dA-D]\)\s+(.+)$')
RE_OPT_3 = re.compile(r'^\s*[iIvVxX]+[.)]\s+(.+)$')

RE_MARKS = re.compile(r'Marks\s*[:=]\s*(\(OR\)|\d+)', re.IGNORECASE)
RE_BTL = re.compile(r'BTL\s*[:=]\s*(\(OR\)|\d+)', re.IGNORECASE)
RE_CO = re.compile(r'(?:CO|Course\s*Outcomes?)\s*[:=]\s*(\(OR\)|[A-Za-z0-9,\s\-]+)', re.IGNORECASE)
RE_EXCEL_TYPE = re.compile(r'Excel\s*Type|excel_type\s*[:=]\s*(.+)$', re.IGNORECASE)

RE_ANSWER = re.compile(r'^(?:Answer\s*Key|Answer|Ans|Correct|KEY)\s*[:=]\s*(.+)$', re.IGNORECASE)
RE_EXPLANATION = re.compile(r'^Explanation\s*[:=]\s*(.+)$', re.IGNORECASE)

RE_CHAPTER = re.compile(r'^\s*(?:Chapter|Section)\s*[:\-]\s*(.+)$', re.IGNORECASE)


def _extract_images_from_paragraph(paragraph, document) -> List[str]:
    """Return list of data URLs for images embedded in this paragraph."""
    out: List[str] = []
    if not paragraph:
        return out

    logger = logging.getLogger('template_api.scan')

    # Primary: look for blip embed relationship ids (common case)
    try:
        blips = paragraph._p.xpath('.//a:blip/@r:embed', namespaces=_NS)
    except Exception:
        blips = []

    for rid in blips or []:
        try:
            part = document.part.related_parts.get(rid)
            if not part:
                continue
            blob = getattr(part, 'blob', None)
            if not blob:
                continue
            content_type = getattr(part, 'content_type', None) or mimetypes.guess_type(getattr(part, 'partname', ''))[0] or 'image/png'
            b64 = base64.b64encode(blob).decode('ascii')
            out.append(f'data:{content_type};base64,{b64}')
        except Exception:
            continue

    # Fallback 1: raw XML r:embed attributes not found by xpath
    try:
        xml = paragraph._p.xml
        for m in re.finditer(r'r:embed\s*=\s*"(rId\d+)"', xml):
            rid = m.group(1)
            try:
                part = document.part.related_parts.get(rid)
                if not part:
                    continue
                blob = getattr(part, 'blob', None)
                if not blob:
                    continue
                content_type = getattr(part, 'content_type', None) or mimetypes.guess_type(getattr(part, 'partname', ''))[0] or 'image/png'
                b64 = base64.b64encode(blob).decode('ascii')
                out.append(f'data:{content_type};base64,{b64}')
            except Exception:
                continue
    except Exception:
        pass

    # Fallback 2: search any attribute values that match a related_parts key
    try:
        for elem in paragraph._p.iter():
            for _, val in elem.attrib.items():
                if val in document.part.related_parts:
                    try:
                        part = document.part.related_parts.get(val)
                        blob = getattr(part, 'blob', None)
                        if not blob:
                            continue
                        content_type = getattr(part, 'content_type', None) or mimetypes.guess_type(getattr(part, 'partname', ''))[0] or 'image/png'
                        b64 = base64.b64encode(blob).decode('ascii')
                        out.append(f'data:{content_type};base64,{b64}')
                    except Exception:
                        continue
    except Exception:
        pass

    # Deduplicate while preserving order
    seen = set()
    deduped: List[str] = []
    for item in out:
        if item not in seen:
            seen.add(item)
            deduped.append(item)

    if logger and deduped:
        try:
            logger.debug('[scan-docx] paragraph images found=%d sample_prefix=%s', len(deduped), (deduped[0][:80] + '...') if isinstance(deduped[0], str) and len(deduped[0])>80 else deduped[0])
        except Exception:
            pass

    return deduped


def _normalize_question_text(s: str) -> str:
    s = (s or '').strip()
    if not s:
        return ''

    m = RE_NUMBERED.match(s)
    if m:
        s = m.group(1).strip()

    m = RE_Q_PREFIX.match(s)
    if m:
        s = m.group(1).strip()

    return re.sub(r'\s+', ' ', s).strip()


def _parse_option_line(s: str) -> Optional[str]:
    s = (s or '').strip()
    if not s:
        return None

    for r in (RE_OPT_1, RE_OPT_2, RE_OPT_3):
        m = r.match(s)
        if m:
            return m.group(1).strip()
    return None


def _co_numbers(co_raw: Optional[str]) -> Optional[str]:
    if not co_raw:
        return None
    nums = sorted({int(x) for x in re.findall(r'\d+', co_raw)})
    if not nums:
        return None
    return ','.join(str(n) for n in nums)


def _parse_answer_value(v: str, options: List[str]) -> Tuple[Optional[str], str]:
    v = (v or '').strip()
    if not v:
        return None, ''

    # Split off inline explanation if present
    expl = ''
    if 'Explanation:' in v:
        parts = v.split('Explanation:', 1)
        v = parts[0].strip()
        expl = parts[1].strip()

    # normalize letter
    m = re.search(r'\(?\s*([A-Da-d])\s*\)?', v)
    if m:
        letter = m.group(1).upper()
        idx = ord(letter) - ord('A')
        if 0 <= idx < len(options):
            return options[idx], expl
        return letter, expl

    return v, expl


def _parse_block(text_lines: List[str], images: List[str], chapter: Optional[str]) -> Optional[Dict[str, Any]]:
    lines = [re.sub(r'\s+', ' ', (x or '').strip()) for x in text_lines if (x or '').strip()]
    if not lines and not images:
        return None

    question_text = _normalize_question_text(lines[0]) if lines else ''
    if not question_text:
        return None

    options: List[str] = []
    answer_text: str = ''
    correct_answer: Optional[str] = None

    marks: Any = None
    btl: Any = None
    co_raw: Any = None
    excel_type: Optional[str] = None

    for ln in lines[1:]:
        # Options
        opt = _parse_option_line(ln)
        if opt is not None:
            options.append(opt)
            continue

        # Answer / Explanation
        m = RE_EXPLANATION.match(ln)
        if m:
            answer_text = (answer_text + ' ' + m.group(1)).strip() if answer_text else m.group(1).strip()
            continue

        m = RE_ANSWER.match(ln)
        if m:
            ca, expl = _parse_answer_value(m.group(1), options)
            correct_answer = ca
            if expl:
                answer_text = (answer_text + ' ' + expl).strip() if answer_text else expl
            continue

        # Metadata
        m = RE_MARKS.search(ln)
        if m:
            marks = m.group(1)
        m = RE_BTL.search(ln)
        if m:
            btl = m.group(1)
        m = RE_CO.search(ln)
        if m:
            co_raw = m.group(1).strip()

        if re.search(r'\bexcel_type\b', ln, re.IGNORECASE) or re.search(r'\bExcelType\b', ln, re.IGNORECASE):
            parts = re.split(r'[:=]', ln, 1)
            if len(parts) == 2:
                excel_type = parts[1].strip()

    # Parenthetical metadata on question line
    qline = lines[0] if lines else ''
    for rx, key in ((RE_MARKS, 'marks'), (RE_BTL, 'btl'), (RE_CO, 'co')):
        m = rx.search(qline)
        if m:
            if key == 'marks':
                marks = m.group(1)
            elif key == 'btl':
                btl = m.group(1)
            else:
                co_raw = m.group(1).strip()

    # OR marker rule
    if str(marks).strip().upper() == '(OR)' and str(btl).strip().upper() == '(OR)' and str(co_raw).strip().upper() == '(OR)':
        # Preserve exact tokens for UI filtering
        marks_out: Any = '(OR)'
        btl_out: Any = '(OR)'
        co_out: Any = '(OR)'
        co_nums: Any = None
    else:
        marks_out = int(marks) if isinstance(marks, str) and marks.isdigit() else (int(marks) if isinstance(marks, int) else None)
        btl_out = int(btl) if isinstance(btl, str) and btl.isdigit() else (int(btl) if isinstance(btl, int) else None)
        co_out = str(co_raw).strip() if co_raw else None
        co_nums = _co_numbers(co_out)

    q: Dict[str, Any] = {
        'question_text': question_text,
        'type': 'objective' if options else None,
        'options': options or None,
        'images': images or None,
        'correct_answer': correct_answer,
        'answer_text': answer_text or '',
        'btl': btl_out,
        'marks': marks_out,
        'chapter': chapter,
        'course_outcomes': co_out,
        'course_outcomes_numbers': co_nums,
        'excel_type': excel_type,
    }

    return q


def _scan_paragraph_mode(document) -> List[Dict[str, Any]]:
    # Deprecated: keep for compatibility, but prefer body-ordered scan
    return _scan_document_in_order(document)


def _scan_table_mode(document) -> Optional[List[Dict[str, Any]]]:
    # Keep for backward compatibility; prefer document-ordered scan.
    # Return None so caller will fall back to paragraph-mode which now uses ordered scan.
    return None


def _scan_document_in_order(document) -> List[Dict[str, Any]]:
    """Iterate through document body elements (paragraphs and tables) in order
    and build a sequential list of question objects. Extract images as data URLs
    and attach them to the nearest question block."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    questions: List[Dict[str, Any]] = []
    current_lines: List[str] = []
    current_images: List[str] = []
    current_chapter: Optional[str] = None

    def flush():
        nonlocal current_lines, current_images
        q = _parse_block(current_lines, current_images, current_chapter)
        if q:
            questions.append(q)
        current_lines = []
        current_images = []

    body = document._element.body
    for child in list(body):
        tag = child.tag.split('}')[-1]
        try:
            if tag == 'p':
                para = Paragraph(child, document)
                txt = (para.text or '').strip()
                imgs = _extract_images_from_paragraph(para, document)

                # Chapter/Section heading
                m = RE_CHAPTER.match(txt)
                if m and not current_lines:
                    current_chapter = m.group(1).strip()
                    continue

                is_blank = (txt == '') and not imgs
                is_numbered = bool(RE_NUMBERED.match(txt)) or bool(RE_Q_PREFIX.match(txt))

                if is_blank:
                    flush()
                    continue

                if is_numbered and current_lines:
                    flush()

                if txt:
                    current_lines.append(txt)
                if imgs:
                    current_images.extend(imgs)

            elif tag == 'tbl':
                # flush any pending paragraph-mode question before handling table rows
                if current_lines:
                    flush()

                table = Table(child, document)
                if not table.rows:
                    continue

                headers = [c.text.strip().lower() for c in table.rows[0].cells]
                if not headers:
                    continue

                # If table doesn't have question-like header, skip
                if not any('question' in h for h in headers):
                    continue

                # build header mapping
                idx = {h: i for i, h in enumerate(headers)}

                def find_col(key: str) -> Optional[int]:
                    for h, i in idx.items():
                        if key in h:
                            return i
                    return None

                c_q = find_col('question')
                c_opt = find_col('option')
                c_marks = find_col('mark')
                c_btl = find_col('btl')
                c_co = find_col('co')
                c_ans = find_col('answer')
                c_img = find_col('image')

                if c_q is None:
                    continue

                for row in table.rows[1:]:
                    cells = row.cells
                    q_images: List[str] = []
                    try:
                        for cp in cells[c_q].paragraphs:
                            q_images.extend(_extract_images_from_paragraph(cp, document))
                    except Exception:
                        q_images = []

                    qtxt = _normalize_question_text(cells[c_q].text)
                    if not qtxt and q_images:
                        # Some templates place only an image in the question cell
                        qtxt = '(Image question)'
                    if not qtxt and not q_images:
                        continue

                    options: List[str] = []
                    images: List[str] = []
                    if q_images:
                        images.extend(q_images)

                    if c_opt is not None:
                        raw_lines = []
                        for cp in cells[c_opt].paragraphs:
                            if cp.text.strip():
                                raw_lines.append(cp.text.strip())
                            images.extend(_extract_images_from_paragraph(cp, document))
                        for ln in raw_lines:
                            opt = _parse_option_line(ln)
                            if opt is not None:
                                options.append(opt)

                    marks_raw = cells[c_marks].text.strip() if c_marks is not None else ''
                    btl_raw = cells[c_btl].text.strip() if c_btl is not None else ''
                    co_raw = cells[c_co].text.strip() if c_co is not None else ''

                    ans_raw = cells[c_ans].text.strip() if c_ans is not None else ''
                    correct_answer = None
                    answer_text = ''
                    if ans_raw:
                        correct_answer, answer_text = _parse_answer_value(ans_raw, options)

                    if c_img is not None:
                        for cp in cells[c_img].paragraphs:
                            images.extend(_extract_images_from_paragraph(cp, document))

                    if str(marks_raw).strip().upper() == '(OR)' and str(btl_raw).strip().upper() == '(OR)' and str(co_raw).strip().upper() == '(OR)':
                        marks_out: Any = '(OR)'
                        btl_out: Any = '(OR)'
                        co_out: Any = '(OR)'
                        co_nums: Any = None
                    else:
                        marks_out = int(marks_raw) if marks_raw.isdigit() else None
                        btl_out = int(btl_raw) if btl_raw.isdigit() else None
                        co_out = co_raw or None
                        co_nums = _co_numbers(co_out)

                    questions.append({
                        'question_text': qtxt,
                        'type': 'objective' if options else None,
                        'options': options or None,
                        'images': images or None,
                        'correct_answer': correct_answer,
                        'answer_text': answer_text or '',
                        'btl': btl_out,
                        'marks': marks_out,
                        'chapter': None,
                        'course_outcomes': co_out,
                        'course_outcomes_numbers': co_nums,
                        'excel_type': None,
                    })

        except Exception:
            # ignore parsing errors for a child and continue
            continue

    # flush any remaining paragraph-mode question
    if current_lines:
        flush()

    return questions


def scan_docx_file(file_obj) -> List[Dict[str, Any]]:
    if Document is None:
        raise RuntimeError('python-docx is not installed')

    document = Document(file_obj)

    table_questions = _scan_table_mode(document)
    if table_questions is not None:
        return table_questions

    # Prefer document-ordered scan which collects paragraphs and tables in flow
    questions = _scan_paragraph_mode(document)

    # Diagnostic logging: count package media parts and images attached
    logger = logging.getLogger('template_api.scan')
    try:
        pkg = getattr(document.part, 'package', None)
        media_parts = []
        if pkg is not None:
            for p in pkg.parts:
                try:
                    partname = getattr(p, 'partname', str(p.partname))
                except Exception:
                    partname = str(p)
                ct = getattr(p, 'content_type', '') or ''
                if ('image' in (ct or '').lower()) or ('/media/' in str(partname).lower()) or str(partname).lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.tiff')):
                    size = None
                    try:
                        blob = getattr(p, 'blob', None)
                        size = len(blob) if blob else None
                    except Exception:
                        size = None
                    media_parts.append({'partname': str(partname), 'content_type': ct, 'size': size})
        total_media = len(media_parts)
    except Exception:
        media_parts = []
        total_media = 0

    try:
        total_questions = len(questions)
        total_images = sum(len(q.get('images') or []) for q in questions)
        per_q = [len(q.get('images') or []) for q in questions[:50]]
        logger.info('scan_docx_file: questions=%d total_images=%d package_media=%d sample_images_per_q=%s', total_questions, total_images, total_media, per_q)
        logger.debug('scan_docx_file: media_parts=%s', media_parts[:20])
    except Exception:
        try:
            logger.info('scan_docx_file: parsed %d questions', len(questions))
        except Exception:
            pass

    return questions
