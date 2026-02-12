"""Test script: ensure images embedded inside a TABLE question cell are extracted.

Run with:
    python backend\manage.py shell -c "exec(open('backend/template_api/scripts/test_scanner_table_image.py').read())"

This creates a docx with a question table and an embedded image inside the
"Answer ALL Questions" cell, then asserts the scanner returns a data URL.
"""

import base64
import os
import tempfile

from docx import Document
from template_api.services.docx_scanner import scan_docx_file

# tiny 1x1 PNG (red)
PNG_B64 = (
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII='
)
PNG_BYTES = base64.b64decode(PNG_B64)


tmpdir = tempfile.mkdtemp()
img_path = os.path.join(tmpdir, 'tiny.png')
doc_path = os.path.join(tmpdir, 'test_table_image.docx')

with open(img_path, 'wb') as f:
    f.write(PNG_BYTES)


doc = Document()

doc.add_paragraph('1. Paragraph question (no image).')

table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
hdr[0].text = 'Q.No.'
hdr[1].text = 'Answer ALL Questions'
hdr[2].text = 'CO'
hdr[3].text = 'BTL'
hdr[4].text = 'Marks'

row = table.add_row().cells
row[0].text = '1'

p = row[1].paragraphs[0]
p.add_run('2. Identify the component shown below:')
p.add_run('\n')
p.add_run().add_picture(img_path)

row[2].text = 'CO1'
row[3].text = 'BTL2'
row[4].text = '2'

doc.save(doc_path)

with open(doc_path, 'rb') as fh:
    questions = scan_docx_file(fh)

print('Parsed questions:', len(questions))
for i, q in enumerate(questions[:5]):
    print('Q', i + 1, 'text=', (q.get('question_text') or '')[:80], 'images=', len(q.get('images') or []))

# Find the table question by its text
q_img = None
for q in questions:
    if 'Identify the component' in (q.get('question_text') or ''):
        q_img = q
        break

assert q_img is not None, 'Expected to find the table question.'
imgs = q_img.get('images') or []
assert isinstance(imgs, list) and imgs, 'Expected at least one extracted image.'
assert isinstance(imgs[0], str) and imgs[0].startswith('data:image/'), 'Expected first image as data URL.'

print('All tests passed. Temporary files in', tmpdir)
