"""
Test script for docx scanner.
Generates a .docx with 15 questions and embeds a tiny PNG into question 3,
then runs scan_docx_file and asserts images are data URLs and count == 15.

Run with:
    python backend\manage.py shell -c "exec(open('backend/template_api/scripts/test_scanner.py').read())"
or directly with the workspace Python environment:
    G:/IDCS-NEW/IDCS-Restart/.venv/Scripts/python.exe backend/template_api/scripts/test_scanner.py
"""
from docx import Document
import base64, tempfile, os
from template_api.services.docx_scanner import scan_docx_file

# tiny 1x1 PNG (red) base64
PNG_B64 = (
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII='
)
PNG_BYTES = base64.b64decode(PNG_B64)

# create temp files
tmpdir = tempfile.mkdtemp()
img_path = os.path.join(tmpdir, 'tiny.png')
with open(img_path, 'wb') as f:
    f.write(PNG_BYTES)

doc_path = os.path.join(tmpdir, 'test_many_questions.docx')
doc = Document()

# create 15 questions
for i in range(1, 16):
    doc.add_paragraph(f"{i}. This is question number {i} which asks something important.")
    # add options for even questions
    if i % 2 == 0:
        doc.add_paragraph('A) Option one')
        doc.add_paragraph('B) Option two')
        doc.add_paragraph('C) Option three')
        doc.add_paragraph('D) Option four')
    # embed image in question 3
    if i == 3:
        doc.add_picture(img_path)

# save
doc.save(doc_path)

# run scanner
with open(doc_path, 'rb') as fh:
    questions = scan_docx_file(fh)

print('Parsed questions:', len(questions))
# find question 3
q3 = None
for q in questions:
    if 'question number 3' in q.get('question_text',''):
        q3 = q
        break

print('Q3 found:', bool(q3))
if q3:
    print('Q3 images:', q3.get('images'))

# simple asserts
assert len(questions) == 15, f'expected 15 questions, got {len(questions)}'
assert q3 is not None, 'question 3 not found'
imgs = q3.get('images') or []
assert isinstance(imgs, list) and len(imgs) >= 1, 'expected embedded image for question 3'
assert isinstance(imgs[0], str) and imgs[0].startswith('data:image/'), 'embedded image not returned as data URL'

print('All tests passed. Temporary files in', tmpdir)
