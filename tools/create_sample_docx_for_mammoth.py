from docx import Document
import base64, os

PNG_B64 = (
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII='
)
PNG_BYTES = base64.b64decode(PNG_B64)

out_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
img_path = os.path.join(out_dir, 'tiny_test_img.png')
doc_path = os.path.join(out_dir, 'sample_for_mammoth.docx')

with open(img_path, 'wb') as f:
    f.write(PNG_BYTES)

doc = Document()
for i in range(1, 8):
    doc.add_paragraph(f"{i}. This is sample question number {i} for testing.")
    if i % 2 == 0:
        doc.add_paragraph('A) Option one')
        doc.add_paragraph('B) Option two')
        doc.add_paragraph('C) Option three')
        doc.add_paragraph('D) Option four')
    if i == 3:
        doc.add_paragraph('Image below:')
        doc.add_picture(img_path)

# add a small table with header mapping and two rows
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Question'
hdr_cells[1].text = 'Options'
hdr_cells[2].text = 'Marks'
hdr_cells[3].text = 'Image'

row = table.add_row().cells
row[0].text = 'Table Q1: What is 2+2?'
row[1].text = 'A) 3\nB) 4\nC) 5\nD) 6'
row[2].text = '4'
row[3].text = ''

row2 = table.add_row().cells
row2[0].text = 'Table Q2: Which is red?'
row2[1].text = 'A) Apple\nB) Banana\nC) Grape\nD) Kiwi'
row2[2].text = '2'
row2[3].text = ''

# Save
if os.path.exists(doc_path):
    os.remove(doc_path)

doc.save(doc_path)
print('Wrote:', doc_path)
print('Wrote image:', img_path)
